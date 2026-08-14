"""Structural checks for the generated Arabic technical report."""

from __future__ import annotations

from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[2]
REPORT = ROOT / "docs" / "report" / "سعيد_بقدونس_هندسة_برمجيات_وذكاء_صنعي_Safe_Autonomous_AI_Agent_VPS.docx"


def main() -> int:
    with ZipFile(REPORT) as package:
        names = set(package.namelist())
        assert "word/document.xml" in names
        media = [name for name in names if name.startswith("word/media/")]
    doc = Document(str(REPORT))
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    required = ["الخلاصة", "الفصل الأول: مواصفات المتطلبات البرمجية (SRS)", "الفصل الثاني: تحليل المتطلبات (SRA)", "الفصل الثالث: تصميم النظام (SD)", "الفصل الرابع: التنفيذ والاختبارات", "الخاتمة التقنية"]
    missing = [item for item in required if item not in headings]
    assert not missing, missing
    assert len(doc.tables) >= 6
    assert len(media) >= 15
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "سعيد بقدونس" in full_text
    assert "[اسم الطالب]" not in full_text
    assert "[اسم_الطالب]" not in full_text
    assert "شبه كود" not in full_text
    assert "pseudocode" not in full_text.lower()
    rtl = 0
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() and paragraph._p.pPr is not None and paragraph._p.pPr.find(qn("w:bidi")) is not None:
            rtl += 1
    assert rtl >= 40, rtl
    print(f"REPORT_STRUCTURAL_VALIDATION=PASS headings={len(headings)} tables={len(doc.tables)} media={len(media)} rtl_paragraphs={rtl}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

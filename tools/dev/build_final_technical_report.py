"""Build the Arabic implementation-focused technical report from the retained template."""

from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
TEMPLATE = ROOT / "1-Report-Template.docx"
OUT = ROOT / "docs" / "report" / "سعيد_بقدونس_هندسة_برمجيات_وذكاء_صنعي_Safe_Autonomous_AI_Agent_VPS.docx"
FIGURES = ROOT / "docs" / "report" / "figures"
BULLET_NUM_ID = None

if not TEMPLATE.exists():
    # The retained final DOCX is a safe local template fallback when the
    # user-provided template is temporarily absent from the repository root.
    TEMPLATE = ROOT / "سعيد_بقدونس_هندسة_برمجيات_وذكاء_صنعي_Safe_Autonomous_AI_Agent_VPS.docx"

if not TEMPLATE.exists():
    TEMPLATE = OUT

BLUE = RGBColor(0x1F, 0x4D, 0x78)
LIGHT_BLUE = "E8EEF5"
GRAY = "F2F4F7"


def set_rtl(paragraph, align=WD_ALIGN_PARAGRAPH.RIGHT):
    paragraph.alignment = align
    ppr = paragraph._p.get_or_add_pPr()
    bidi = ppr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        ppr.append(bidi)
    bidi.set(qn("w:val"), "1")


def set_ltr(paragraph, align=WD_ALIGN_PARAGRAPH.LEFT):
    """Configure an English paragraph without inheriting the document RTL flow."""
    paragraph.alignment = align
    ppr = paragraph._p.get_or_add_pPr()
    bidi = ppr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        ppr.append(bidi)
    bidi.set(qn("w:val"), "0")


def set_font(run, name="Arial", size=12, bold=False, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def set_run_direction(run, *, rtl: bool):
    """Make mixed Arabic/Latin runs deterministic in Word's bidi layout."""
    rpr = run._element.get_or_add_rPr()
    direction = rpr.find(qn("w:rtl"))
    if direction is None:
        direction = OxmlElement("w:rtl")
        rpr.append(direction)
    direction.set(qn("w:val"), "1" if rtl else "0")
    lang = rpr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        rpr.append(lang)
    lang.set(qn("w:val"), "ar-SA" if rtl else "en-US")


def add_mixed_runs(paragraph, text, *, size=12, bold=False, color=None):
    """Add text as explicit RTL Arabic and LTR Latin runs without reordering."""
    arabic = r"[\u0600-\u08ff\uFB50-\uFDFF\uFE70-\uFEFF]"
    parts = [part for part in re.split(rf"({arabic}+)", str(text)) if part]
    for part in parts:
        is_rtl = bool(re.search(arabic, part))
        run = paragraph.add_run(part)
        set_font(run, size=size, bold=bold, color=color)
        set_run_direction(run, rtl=is_rtl)


def shade(cell, fill: str):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcpr = tc.get_or_add_tcPr()
    tc_mar = tcpr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tcpr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    table.autofit = False
    tbl = table._tbl
    tblpr = tbl.tblPr
    tblw = tblpr.find(qn("w:tblW"))
    if tblw is None:
        tblw = OxmlElement("w:tblW")
        tblpr.append(tblw)
    tblw.set(qn("w:w"), str(sum(widths)))
    tblw.set(qn("w:type"), "dxa")
    ind = tblpr.find(qn("w:tblInd"))
    if ind is None:
        ind = OxmlElement("w:tblInd")
        tblpr.append(ind)
    ind.set(qn("w:w"), "120")
    ind.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tcpr = cell._tc.get_or_add_tcPr()
            tcw = tcpr.find(qn("w:tcW"))
            if tcw is None:
                tcw = OxmlElement("w:tcW")
                tcpr.append(tcw)
            tcw.set(qn("w:w"), str(widths[i]))
            tcw.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def mark_header_row(row):
    trpr = row._tr.get_or_add_trPr()
    header = trpr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        trpr.append(header)
    header.set(qn("w:val"), "true")


def add_page_field(paragraph, field="PAGE"):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {field} "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_font(run, size=10, color=RGBColor(0x66, 0x66, 0x66))


def add_toc_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "افتح الحقل في Word لتحديث فهرس المحتويات"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_font(run, size=12)


def add_text(doc, text, *, bold=False, size=12, align=WD_ALIGN_PARAGRAPH.RIGHT, style=None):
    p = doc.add_paragraph(style=style)
    set_rtl(p, align)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    add_mixed_runs(p, text, size=size, bold=bold)
    return p


def add_ltr_text(doc, text, *, bold=False, size=12, align=WD_ALIGN_PARAGRAPH.LEFT):
    """Add a genuinely left-to-right paragraph for the English abstract."""
    p = doc.add_paragraph()
    set_ltr(p, align)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_font(run, name="Arial", size=size, bold=bold)
    set_run_direction(run, rtl=False)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Paragraph")
    set_rtl(p)
    p.paragraph_format.space_after = Pt(3)
    ppr = p._p.get_or_add_pPr()
    numpr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numid = OxmlElement("w:numId")
    numid.set(qn("w:val"), str(BULLET_NUM_ID))
    numpr.extend([ilvl, numid])
    ppr.append(numpr)
    add_mixed_runs(p, text, size=11)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    set_rtl(p)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    add_mixed_runs(p, text, size={1: 18, 2: 15, 3: 13}.get(level, 12), bold=True, color=BLUE)
    return p


def add_table(doc, headers: list[str], rows: Iterable[Iterable[str]], widths=None, caption=None):
    if caption:
        add_text(doc, caption, bold=True, size=10)
    rows = [list(map(str, row)) for row in rows]
    table = doc.add_table(rows=1, cols=len(headers))
    mark_header_row(table.rows[0])
    widths = widths or [9360 // len(headers)] * len(headers)
    set_table_geometry(table, widths)
    for i, value in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade(cell, LIGHT_BLUE)
        cell.text = ""
        p = cell.paragraphs[0]
        set_rtl(p, WD_ALIGN_PARAGRAPH.CENTER)
        add_mixed_runs(p, value, size=9, bold=True, color=BLUE)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            set_rtl(p, WD_ALIGN_PARAGRAPH.RIGHT)
            add_mixed_runs(p, value, size=8.5)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_figure(doc, number, title, filename):
    p = doc.add_paragraph()
    set_rtl(p, WD_ALIGN_PARAGRAPH.CENTER)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    inline = run.add_picture(str(FIGURES / filename), width=Inches(6.0))
    inline._inline.docPr.set("descr", title)
    inline._inline.docPr.set("title", title)
    cap = doc.add_paragraph()
    set_rtl(cap, WD_ALIGN_PARAGRAPH.CENTER)
    add_mixed_runs(cap, f"الشكل {number}: {title}", size=10, bold=True, color=BLUE)
    cap.paragraph_format.space_after = Pt(8)


def configure_document(doc):
    global BULLET_NUM_ID
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    BULLET_NUM_ID = max(num_ids, default=0) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), "bullet")
    text = OxmlElement("w:lvlText")
    text.set(qn("w:val"), "•")
    jc = OxmlElement("w:lvlJc")
    jc.set(qn("w:val"), "left")
    ppr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    ppr.append(ind)
    level.extend([start, fmt, text, jc, ppr])
    abstract.append(level)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(BULLET_NUM_ID))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
        section.header_distance = Inches(0.49)
        section.footer_distance = Inches(0.49)
        footer = section.footer.paragraphs[0]
        footer.text = ""
        set_rtl(footer, WD_ALIGN_PARAGRAPH.CENTER)
        r = footer.add_run("التقرير التقني - ")
        set_font(r, size=9, color=RGBColor(0x66, 0x66, 0x66))
        add_page_field(footer)
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:cs"), "Arial")
    normal.font.size = Pt(12)
    for name, size in (("Heading 1", 18), ("Heading 2", 15), ("Heading 3", 13)):
        style = doc.styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style._element.rPr.rFonts.set(qn("w:cs"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = BLUE


def clear_body(doc):
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def cover(doc):
    for text, size, bold in (
        ("الجمهورية العربية السورية", 16, True),
        ("المعهد العالي للعلوم التطبيقية والتكنولوجيا", 15, True),
        ("هندسة برمجيات وذكاء صنعي", 14, False),
    ):
        add_text(doc, text, bold=bold, size=size, align=WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(3):
        doc.add_paragraph()
    add_text(doc, "تقرير تقني نهائي", bold=True, size=22, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, "وكيل ذكاء صنعي آمن ومستقل لإدارة الخوادم الافتراضية الخاصة", bold=True, size=20, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, "Safe Autonomous AI Agent for VPS Management", bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(4):
        doc.add_paragraph()
    add_text(doc, "إعداد: سعيد بقدونس", size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, "إشراف: المهندس محمد بشار الدسوقي", size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, "السنة الرابعة - هندسة برمجيات وذكاء صنعي", size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, "2026", size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()


def abstract_and_front(doc):
    add_heading(doc, "الخلاصة", 1)
    add_text(doc, "يتضمن هذا التقرير شرحاً للمشروع المنفذ لإدارة خوادم Linux باستخدام الذكاء الصنعي. يراقب النظام الخادم، ويحفظ التقارير، ويحللها، ويطلب فحصاً إضافياً عند الحاجة. كما يمكنه اقتراح معالجة المشكلة، لكن تنفيذها يمر دائماً عبر التحقق والموافقة والقيود الموضوعة في النظام.")
    add_text(doc, "يعتمد المشروع على Claude Code للإشراف، وعلى Ollama لتحليل النصوص، بينما تنفذ Python وقاعدة PostgreSQL العمليات المهمة. أظهرت الاختبارات سلامة المصادقة والصلاحيات، وحدود الأدوات، وحفظ الأدلة، ومنع التكرار، والتعامل مع التزامن والأخطاء. وقد فصلنا في هذا التقرير بين نتائج الاختبارات المحلية ونتائج القبول التي تحتاج إلى بيئة تشغيل حقيقية.")
    add_heading(doc, "ملخص تنفيذي", 1)
    add_text(doc, "يقدم هذا التقرير وصفاً للمشروع المنفذ لإدارة خوادم Linux باستخدام الذكاء الصنعي. يراقب النظام الخادم ويحفظ التقارير ويحللها، ثم يجمع الأدلة عند الحاجة. ويمكنه اقتراح المعالجة، لكن التنفيذ يمر دائماً عبر التحقق والموافقة والقيود المسجلة. أظهرت الاختبارات المحلية سلامة الصلاحيات وحدود الأدوات وحفظ الأدلة ومنع التكرار والتعامل مع التزامن والأخطاء، بينما بقي قبول البيئات الحقيقية منفصلاً عن نتائج الاختبارات المحلية.")
    doc.add_page_break()
    add_heading(doc, "Abstract", 1)
    add_ltr_text(doc, "This report presents the analysis, design, implementation, and testing of a safe autonomous AI agent for Linux VPS management. The system collects bounded read-only observations, stores monitoring reports, analyses them with an Ollama language model, and requests additional evidence through registered specialist capabilities when the initial diagnosis is insufficient. Any remediation proposal is represented as a plan and passes policy checks, sandbox validation or human approval, authorization, verification, and audit before it can affect a server.")
    add_ltr_text(doc, "The implementation keeps the execution authority in Python services and PostgreSQL rather than in the language model. Claude Code supervises the conversation through a typed MCP boundary, while SSH is restricted by known-host verification and named operations. The report distinguishes deterministic local test results from real-runtime acceptance that depends on Ollama, PostgreSQL, SSH, and an isolated target environment.")
    doc.add_page_break()
    add_heading(doc, "فهرس المحتويات", 1)
    toc_entries = [
        "الخلاصة ........................................................................ 2",
        "Abstract ........................................................................ 3",
        "مقدمة عامة ...................................................................... 9",
        "الفصل الأول: التعريف بالمشروع والمتطلبات .............................. 9",
        "1.1 هدف المشروع ونطاقه ..................................................... 9",
        "1.3 المتطلبات الوظيفية ....................................................... 10",
        "1.5 المتطلبات غير الوظيفية والقبول ....................................... 13",
        "1.6 مصفوفة التتبع وخطة التنفيذ والمخاطر ................................ 14",
        "الفصل الثالث: الدراسة التحليلية ............................................ 16",
        "3.1 الممثلون وحالات الاستخدام ............................................. 16",
        "3.1.4 بطاقات حالات الاستخدام ............................................... 18",
        "3.2 دورة الحادثة وسيناريوهات الفشل ....................................... 24",
        "الفصل الرابع: الدراسة التصميمية ........................................... 27",
        "4.1 مخطط النظام وسياقه ...................................................... 27",
        "4.6 المتخصصون وEvidence .................................................... 31",
        "4.8 المعالجة الذاتية والحجز .................................................. 34",
        "4.14 تصميم الاتساق والتزامن .................................................. 39",
        "4.19 الأدوات والبيئات المستخدمة ............................................ 41",
        "الفصل الخامس: التنفيذ والاختبارات ......................................... 43",
        "5.1 البيئة والتقنيات وسبب الاختيار .......................................... 43",
        "5.6 دورة المعالجة الذاتية ...................................................... 44",
        "5.8 نتائج الاختبار الحالية ..................................................... 45",
        "الفصل السادس: الواجهات والاختبار .......................................... 46",
        "6.2 الواجهة الرئيسية ومتابعة الخوادم ....................................... 46",
        "6.4 واجهة التحقيق والأدلة ...................................................... 47",
        "6.5 واجهة المعالجة والموافقة .................................................. 48",
        "6.8 خطة الاختبار ونتائج القبول .............................................. 50",
        "الخاتمة والآفاق المستقبلية .................................................... 52",
        "الملاحق ............................................................................. 53",
    ]
    for entry in toc_entries:
        add_text(doc, entry, size=10)
    add_heading(doc, "قائمة الأشكال", 1)
    for i, title in enumerate(["مخطط حالات الاستخدام", "مخطط تسلسل المراقبة والتحليل", "مخطط نشاط المعالجة الخاضعة للإشراف", "مخطط النظام العام", "سياق النظام", "طبقات المكونات", "تدفق Claude وOllama وMCP وPython", "المراقبة والتحليل والتحقيق", "تنسيق المتخصصين", "المعالجة الخاضعة للإشراف", "التحقق في Sandbox", "التنفيذ الذاتي", "السياسة والحجز وعدم التكرار", "قاطع الدارة والاستعادة", "مصادقة Admin والصلاحيات", "البنية الفيزيائية", "علاقات قاعدة البيانات", "المكونات الأساسية", "حالات استخدام المشغل", "حالات استخدام مدير النظام", "تسلسل التحقيق المتخصص", "تسلسل المعالجة الخاضعة للإشراف", "تسلسل التنفيذ الذاتي", "تسلسل موافقة المدير والتدقيق", "واجهة لوحة الإدارة", "واجهة التحقيق والأدلة", "واجهات المعالجة والسياسات"], 1):
        p = add_text(doc, f"الشكل {i}: {title}", size=9)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.0
    add_heading(doc, "قائمة الجداول", 1)
    for i, title in enumerate(["حالة المتطلبات", "المتطلبات غير الوظيفية ومعايير قبولها", "مصفوفة تتبع المتطلبات وحالات الاستخدام", "مراحل تنفيذ المشروع", "مخاطر المشروع وطرق الحد منها", "عينة حالات الاستخدام", "الوصف التفصيلي لحالات الاستخدام الرئيسية", "قرارات المعالجة الذاتية", "سيناريوهات الفشل", "المبادئ والأنماط التصميمية", "الأدوات والبيئات وحدود استخدامها", "التقنيات الرئيسية", "النتائج الحالية", "صفحات الواجهة ووظائفها", "مصفوفة أنواع الاختبار وحدودها", "جرد الجداول", "دلالة مجموعات قاعدة البيانات", "عينة حالات الاستخدام الأساسية", "مصفوفة الاختبارات ونتائجها"], 1):
        p = add_text(doc, f"الجدول {i}: {title}", size=9)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.0
    doc.add_page_break()
    add_heading(doc, "المصطلحات والاختصارات", 1)
    add_table(doc, ["الاختصار", "الاسم باللغة الإنكليزية", "المعنى في المشروع"], [["API", "Application Programming Interface", "واجهة طلبات منظمة بين العميل والخدمات."], ["MCP", "Model Context Protocol", "حد الأدوات typed الذي يراه Claude."], ["SSH", "Secure Shell", "اتصال موثق بالخادم عبر known_hosts."], ["VPS", "Virtual Private Server", "الخادم الافتراضي الذي تتم مراقبته."], ["LLM", "Large Language Model", "نموذج لغوي يساعد في تحليل النصوص."], ["RAG", "Retrieval-Augmented Generation", "استرجاع سياق من تقارير ووثائق مفهرسة."], ["RBAC", "Role-Based Access Control", "التحكم بالوصول المبني على الدور."], ["CSRF", "Cross-Site Request Forgery", "حماية طلبات التعديل في واجهة الإدارة."], ["SQL", "Structured Query Language", "لغة الاستعلام المستخدمة في قاعدة البيانات."], ["ORM", "Object-Relational Mapping", "ربط نماذج Python بجداول PostgreSQL."], ["ERD", "Entity Relationship Diagram", "مخطط علاقات الكيانات والجداول."], ["REST", "Representational State Transfer", "أسلوب تنظيم واجهات HTTP."], ["UI", "User Interface", "واجهة Admin التي يستخدمها المشغل والمدير."], ["DB", "Database", "مخزن البيانات التشغيلية والأدلة والتدقيق."], ["Evidence", "Operational Evidence", "دليل مرتبط بحادثة وخادم ومصدر ووقت."], ["Sandbox", "Isolated Validation Environment", "بيئة تحقق معزولة قبل الإجراء."], ["TTL", "Time To Live", "مدة صلاحية الحجز أو التصريح."], ["NFR", "Non-Functional Requirement", "صفة جودة قابلة للفحص وليست وظيفة جديدة."]], [1300, 3300, 4760])
    doc.add_page_break()


def chapter1(doc):
    add_heading(doc, "الفصل الأول: التعريف بالمشروع والمتطلبات", 1)
    add_heading(doc, "1.1 هدف المشروع ونطاقه", 2)
    add_text(doc, "يهدف المشروع إلى مساعدة المشغل على اكتشاف أعطال الخوادم الافتراضية بسرعة، وفهم سبب العطل، واختيار المعالجة المناسبة. يراقب النظام الموارد والخدمات والسجلات، ويحفظ النتائج، ويحللها، ثم يجمع الأدلة عند الحاجة. ولا يسمح للنموذج اللغوي بتنفيذ أوامر مباشرة على الخادم.")
    add_text(doc, "يشمل الإصدار الحالي المراقبة والتحليل والتحقيق والمتابعة والمعالجة الخاضعة للإشراف وبعض حواجز التنفيذ الذاتي. أما الإشعارات الاجتماعية والتنبؤ الإنتاجي وتعديل كود التطبيق تلقائياً فليست ضمن الإصدار الحالي.")
    add_heading(doc, "1.2 أصحاب المصلحة والقيود", 2)
    for x in ["المشغّل: يراقب الخوادم ويتعامل مع نتائج المعالجة الخاضعة للإشراف.", "المطور/الموافق: يقرر الإجراءات الخطرة أو الحساسة عبر Admin المحلي.", "المراجع: يحتاج إلى تتبع المتطلبات والأدلة والاختبارات.", "النظام: ينفذ القراءة والتحقق وفق حدود Python وPostgreSQL وSSH.", "القيود: PostgreSQL/Ollama/Claude/WSL2 وهدف آمن مطلوب للقبول الحي؛ لا توجد أسرار في الوثائق."]:
        add_bullet(doc, x)
    add_heading(doc, "1.3 المتطلبات الوظيفية", 2)
    add_text(doc, "توجد المتطلبات الوظيفية كاملة في docs/requirements/functional-requirements.md. يلخص الجدول التالي أهم الوظائف وحالتها الحالية:")
    add_table(doc, ["الفئة", "الحالة الحالية", "الدليل"], [["المراقبة والتحليل والتحقيق", "منفذة", "اختبارات الوحدات والتكامل وC.14.12"], ["الأدلة والمتخصصون", "منفذة", "اختبارات الملكية والتجميع والتخزين"], ["المعالجة الخاضعة للإشراف", "منفذة", "اختبارات Phase 5 وسجل القبول"], ["بيئة الاختبار المعزولة", "جزئية", "العقود والاختبارات موجودة، ودليل القبول يحتاج توحيداً"], ["المعالجة الذاتية", "منفذة، والقبول الحي مؤجل", "اختبارات السياسة والحجز وقاطع الدارة"], ["الإشعار الاجتماعي", "مؤجل", "لا يوجد موصل Telegram أو قناة اجتماعية"], ["تحديد موضع خطأ الكود", "جزئية", "التشخيص موجود دون مسار مصدر مخصص"]], [2600, 2800, 3760], "الجدول 1: ملخص المتطلبات الوظيفية")
    add_heading(doc, "1.3.1 وظائف المشغل", 3)
    for x in [
        "FR-01: اختيار خادم وملف مراقبة محفوظ وتشغيل مجموعة أوامر القراءة المحددة.",
        "FR-02: عرض تقرير المراقبة مع حالة كل أمر ووقت التنفيذ وحجم المخرجات.",
        "FR-03: طلب تحليل تقرير سابق وربط التحليل بالتقرير ومصادره.",
        "FR-04: طلب تحقيق إضافي عند نقص المعلومات أو اختلاف نتائج التحليل.",
        "FR-05: عرض الأدلة ومصدر كل دليل والمتخصص الذي جمعه والحادثة المرتبطة به.",
        "FR-06: مراجعة التشخيص وخطة المعالجة والخطر والهدف قبل اتخاذ القرار.",
        "FR-07: الموافقة على الخطة أو رفضها أو إعادتها للمراجعة مع تسجيل السبب.",
        "FR-08: تشغيل معالجة خاضعة للإشراف ثم عرض نتيجة التحقق والاستعادة عند الحاجة.",
        "FR-09: متابعة حالة التصريح والحجز والتنفيذ وسجل التدقيق دون كشف الرموز السرية.",
    ]:
        add_bullet(doc, x)
    add_heading(doc, "1.3.2 وظائف مدير النظام", 3)
    for x in [
        "FR-10: تسجيل الدخول والخروج وإدارة الجلسة وفق الدور الممنوح.",
        "FR-11: إدارة الخوادم وملفات المراقبة والمتخصصين ومصادر المعرفة.",
        "FR-12: مراجعة سياسات التنفيذ الذاتي وتفعيلها أو إيقافها ضمن الصلاحية.",
        "FR-13: مراجعة قرارات السياسة والتفويضات والحجوزات ومحاولات الاسترداد.",
        "FR-14: البحث في سجل التدقيق ومراجعة محاولات الرفض والفشل والتكرار.",
        "FR-15: عرض حالة الخدمات وقاعدة البيانات وبيئة التنفيذ من خلال صفحات الإدارة.",
    ]:
        add_bullet(doc, x)
    add_heading(doc, "1.3.3 وظائف النظام الداخلية", 3)
    for x in [
        "FR-16: تنفيذ أوامر المراقبة المسجلة فقط مع مهلة وحجم مخرجات محدودين.",
        "FR-17: حفظ التقارير والتحليلات والأدلة والخطط والقرارات في قاعدة البيانات.",
        "FR-18: تقييم السياسة قبل أي تنفيذ ومنع المسار عند نقص الدليل أو اختلاف السياق.",
        "FR-19: إصدار تصريح لمرة واحدة وحجز قصير وربط الحفظ النهائي بمالك العملية.",
        "FR-20: منع إعادة تنفيذ العملية نفسها أو فقدان تحديث عامل آخر عند التزامن.",
        "FR-21: التحقق من النتيجة بعد الإجراء وتسجيل الفشل والاستعادة عند عدم تحققها.",
        "FR-22: استعادة الحجز المنتهي بشروط واضحة أو التوقف عند غموض حالة التنفيذ.",
    ]:
        add_bullet(doc, x)
    add_heading(doc, "1.4 تدقيق الوظائف 1-9 في دفتر الشروط", 2)
    add_text(doc, "تم تنفيذ الوظائف من 1 إلى 6 والوظيفة 9، مع وجود بدائل محلية لبعضها. أما الوظيفة 7 فما زالت جزئية، لأن النظام يحفظ التشخيص والأدلة لكنه لا يحدد مكان الخطأ داخل كود التطبيق بشكل كامل. والوظيفة 8 مؤجلة، إذ لا توجد حالياً قناة Telegram أو قناة اجتماعية. وتحتاج نتيجتا Phase 6 وPhase 7 إلى توثيق قبول حي مستقل قبل إعلان الإغلاق النهائي.")
    add_heading(doc, "1.5 المتطلبات غير الوظيفية والقبول", 2)
    add_text(doc, "لا تصف المتطلبات غير الوظيفية وظيفة جديدة، بل تحدد مستوى الجودة والقيود التي يجب أن تبقى صحيحة أثناء تنفيذ الوظائف. لذلك صيغت هنا على شكل صفات قابلة للفحص، وربط كل صفة بطريقة تحقق واضحة بدلاً من تحويل أعداد الجداول أو الأدوات إلى متطلبات جودة.")
    add_table(doc, ["الصفة", "المتطلب", "طريقة التحقق"], [
        ["الأمن", "لا يصل المستخدم أو النموذج إلى تنفيذ غير مسجل، وتطبق الجلسات وRBAC وCSRF وknown_hosts.", "اختبارات الصلاحيات والطلبات ورفض shell الحر وعدم كشف الأسرار"],
        ["الموثوقية", "عند تعطل قاعدة البيانات أو النموذج أو SSH يتوقف المسار الآمن ويحفظ سبب الفشل دون قرار ناقص.", "اختبارات الفشل، النتائج غير المكتملة، وقاطع الدارة"],
        ["الأداء", "تحدد أوامر المراقبة مهلة وحجم مخرجات، ولا تبقى معاملة قاعدة البيانات مفتوحة أثناء اتصال Ollama أو SSH.", "اختبارات المهلة ومراجعة مسار الحجز والتنفيذ"],
        ["التزامن وعدم التكرار", "لا يؤدي إرسال الطلب مرتين أو تشغيل عاملين إلى تنفيذ مكرر أو فقدان تحديث سابق.", "بصمة العملية، تصريح واحد، حجز قصير، owner token وتحديث ذري"],
        ["قابلية الاستخدام", "يعرض Admin الحالة والقرار والسبب والأدلة والوقت بصيغة يفهمها المشغل.", "اختبارات القوالب والاستجابات ومراجعة الصفحات"],
        ["قابلية الصيانة", "تفصل العقود والسياسات عن القدرات والبنية التحتية والواجهات، وتستخدم composition واحداً.", "فحص بنية الحزم، الاستيرادات، واختبارات العقود"],
        ["قابلية التوسع", "يمكن إضافة خادم أو متخصص أو سياسة دون إنشاء فرع تنفيذي خاص يلتف على القيود.", "تعريفات قاعدة البيانات واختبارات أكثر من حالة"],
        ["القابلية للتتبع", "ترتبط كل نتيجة بمصدرها ومالكها وزمنها والحادثة التي تنتمي إليها.", "فحص Evidence وسجل التدقيق ومصفوفة التتبع"],
        ["قابلية الاختبار", "تعمل الاختبارات المحلية دون Ollama أو SSH حقيقي، وتبقى اختبارات القبول الخارجي opt-in.", "pytest وTestClient وفصل real_runtime"],
    ], [1900, 4700, 2760], "الجدول 2: المتطلبات غير الوظيفية ومعايير قبولها")
    add_text(doc, "تعد هذه المتطلبات محققة عندما تظهر في السلوك والاختبارات، لا لمجرد وجود ملف أو مكوّن. أما عدد الجداول وأدوات MCP والمسارات فيعرض لاحقاً كمؤشر جرد للتنفيذ، وليس بديلاً عن متطلبات الأمن والموثوقية والأداء.")
    add_heading(doc, "1.6 مصفوفة التتبع", 2)
    add_text(doc, "تربط مصفوفة التتبع في docs/requirements/traceability-matrix.md كل متطلب بالمكوّن المسؤول والاختبار والدليل. لذلك لا نعتبر المتطلب ناجحاً اعتماداً على تقرير قديم، بل نحتاج إلى دليل يمكن التحقق منه.")
    add_table(doc, ["المتطلب", "حالة الاستخدام", "المكوّن المسؤول", "دليل التحقق"], [
        ["FR-01 إلى FR-03", "UC-001، UC-005", "MonitoringService وAnalysisOrchestrator", "اختبارات المراقبة والتحليل وحفظ التقرير"],
        ["FR-04 إلى FR-06", "UC-009، UC-010، UC-011", "InvestigationRouter وSpecialistInvestigationLoop", "اختبارات السياسة وتجميع Evidence والملكية"],
        ["FR-07 إلى FR-09", "UC-017، UC-018", "Admin UI وRemediationService", "اختبارات الموافقة والجلسة والتدقيق"],
        ["FR-10 إلى FR-15", "UC-025، UC-026", "AdminAuthService وواجهات الإدارة", "اختبارات RBAC وCSRF والقوالب"],
        ["FR-16 إلى FR-18", "UC-001، UC-009، UC-019", "Capabilities وPolicyEvaluator", "اختبارات حدود الأدوات والفشل المغلق"],
        ["FR-19 إلى FR-22", "UC-019، UC-022", "Authorization وReservation وExecution", "اختبارات التصريح والحجز والتزامن والاستعادة"],
        ["NFR-01 إلى NFR-09", "جميع الحالات", "app/core وcomposition والواجهات", "الاختبارات المحلية وفحص البنية والقبول المنفصل"],
    ], [1800, 2200, 3000, 2860], "الجدول 3: مصفوفة تتبع المتطلبات وحالات الاستخدام")
    add_heading(doc, "1.7 المتطلبات التفصيلية للمراقبة والتحليل", 2)
    add_text(doc, "يجب أن يختار المشغل ملف مراقبة محفوظاً لكل خادم. يحتوي الملف على أوامر القراءة المسموحة والمهلة وحجم المخرجات. وتحفظ النتيجة مع اسم الخادم ووقت التنفيذ واسم الأمر وحالته، حتى نعرف مصدر كل معلومة في التقرير.")
    add_text(doc, "يبحث النظام أولاً عن تحليل سابق يمكن استخدامه من جديد. وإذا احتاج إلى تحليل جديد، يستخدم التقارير والوثائق المناسبة. ويسجل مصدر كل معلومة، ويميز بين التقرير الحديث ووثيقة المعرفة العامة. أما إذا تعطل مزود التحليل أو كانت البيانات ناقصة، فيحفظ النظام سبب الفشل ولا يعتبره نجاحاً.")
    add_heading(doc, "1.8 المتطلبات التفصيلية للتحقيق والأدلة", 2)
    add_text(doc, "يعتمد التحقيق على متخصّصين محفوظين في قاعدة البيانات. لكل متخصّص اسم وهدف وأدوات مسموحة وميزانية محددة. ولا تكفي إجابة النموذج وحدها، بل تحفظ الملاحظات والأدلة ومصدر كل ملاحظة والخادم الذي جاءت منه.")
    add_text(doc, "يرفض النظام الدليل الذي لا يخص الحادثة أو الخادم المطلوب. وإذا اختلف متخصّصان، يظهر الاختلاف للمشغل بدلاً من إخفائه. وبعد جمع الأدلة تتم مقارنة النتائج، ثم يصدر التشخيص النهائي.")
    add_heading(doc, "1.9 المتطلبات التفصيلية للمعالجة الآمنة", 2)
    add_text(doc, "لا ينفذ النظام المعالجة لمجرد أن التحليل اقترحها. تحتوي كل خطة على الإجراء والخادم والهدف، وترتبط بالموافقة أو بنتيجة الاختبار المعزول. ويعيد النظام فحص هذه المعلومات عند التنفيذ، لأن الخطة أو السياسة قد تتغير بعد إنشائها.")
    add_text(doc, "يمنع النظام تكرار العملية إذا أُرسل الطلب مرتين أو عمل عاملان في الوقت نفسه. ويستخدم لذلك مفتاحاً خاصاً للعملية وحجزاً مؤقتاً ومالكاً واضحاً للمعالجة. وإذا حدث شك أو فشل، يتوقف النظام ويحفظ السبب بدلاً من التخمين.")
    add_heading(doc, "1.10 حدود الإصدار الحالي ومعايير الإغلاق", 2)
    add_text(doc, "يكون الإصدار الحالي ناجحاً عندما تعمل المراقبة والتحليل والتحقيق والمعالجة المشرفة ضمن الاختبارات، وعندما تكون قاعدة البيانات والأدوات والواجهات قابلة للفحص. أما الإغلاق النهائي فيحتاج إلى تشغيل بيئة Ollama وSSH وSandbox وتسجيل نتيجة قبول حقيقية لكل مرحلة.")
    add_heading(doc, "1.11 خطة تنفيذ المشروع", 2)
    add_text(doc, "اتبع تنفيذ المشروع مراحل متتابعة تبدأ بفهم المشكلة وتحديد الحدود، ثم بناء العقود وقاعدة البيانات، ثم إضافة المراقبة والتحليل والتحقيق، ثم المعالجة المشرفة والتنفيذ الذاتي، وأخيراً الواجهات والاختبارات والقبول. يسمح هذا الترتيب بإثبات القراءة والحفظ قبل السماح بأي تغيير على الخادم.")
    add_table(doc, ["المرحلة", "العمل المنجز", "المخرج"], [
        ["تحليل المشكلة", "تحديد الخادم والمشغل والمخاطر والوظائف", "نطاق ومتطلبات"],
        ["بناء الأساس", "العقود والسياسات وcomposition وقاعدة البيانات", "هيكل قابل للاختبار"],
        ["المراقبة والتحليل", "أوامر القراءة والتقارير والتحليل والمصادر", "تقرير قابل للمراجعة"],
        ["التحقيق", "المتخصصون وحلقة الأدلة والتشخيص", "Evidence مستمر"],
        ["المعالجة", "الخطة والموافقة وSandbox والتحقق والاستعادة", "تنفيذ مشرف"],
        ["التنفيذ الذاتي", "السياسة والتصريح والحجز والمالك وقاطع الدارة", "تنفيذ مقيد"],
        ["الاختبار والقبول", "اختبارات محلية ثم بيئة تشغيل حقيقية اختيارية", "نتيجة موثقة وحدود معلنة"],
    ], [1900, 4700, 2760], "الجدول 4: مراحل تنفيذ المشروع")
    add_heading(doc, "1.12 مخاطر المشروع والحدود العملية", 2)
    add_text(doc, "ترتبط المخاطر الأساسية بأن النظام يتعامل مع خوادم حقيقية وبنتائج قد تكون ناقصة أو قديمة. لذلك لا يعالجها المشروع بزيادة استقلالية النموذج، بل بتقييد المدخلات، والتحقق من السياق، وحفظ الأدلة، والموافقة، والفشل المغلق، وفصل المعاملات القصيرة عن الاتصالات الخارجية.")
    add_table(doc, ["الخطر", "الأثر المحتمل", "إجراء الحد منه"], [
        ["تشخيص ناقص", "اقتراح معالجة غير مناسبة", "طلب Evidence إضافي ورفض المعالجة عند النقص"],
        ["تغير الخادم بعد إنشاء الخطة", "تنفيذ خطة قديمة", "بصمة الخطة وفحص الهدف قبل التنفيذ"],
        ["عاملان متزامنان", "تنفيذ مزدوج أو فقدان تحديث", "مفتاح فريد وحجز وowner token ودمج ذري"],
        ["انقطاع SSH", "نتيجة تنفيذ غير واضحة", "تحقق مستقل وتوقف عند الغموض وسجل تدقيق"],
        ["تسرب سر أو صلاحية", "وصول غير مصرح", "RBAC وCSRF وإسقاطات صريحة وknown_hosts"],
        ["تعطل خدمة خارجية", "تحليل أو قبول غير مكتمل", "حفظ الفشل وفصل الاختبارات الحقيقية عن المحلية"],
    ], [2200, 3300, 3860], "الجدول 5: مخاطر المشروع وطرق الحد منها")


def chapter2_theory(doc):
    """Write the theory and reference chapter in the style used by prior reports."""
    add_heading(doc, "الفصل الثاني: الدراسة النظرية والمرجعية", 1)
    add_heading(doc, "2.1 مقدمة الدراسة النظرية", 2)
    add_text(doc, "يتناول هذا الفصل المفاهيم التي يعتمد عليها المشروع قبل الانتقال إلى تحليل المتطلبات وتصميم النظام. والهدف منه شرح المشكلة والمصطلحات والأدوات التي تجعل إدارة الخوادم أكثر انتظاماً، مع توضيح سبب اختيار التقنيات المستخدمة في النسخة الحالية.")
    add_heading(doc, "2.2 الخوادم الافتراضية وإدارة الموارد", 2)
    add_text(doc, "الخادم الافتراضي هو بيئة تشغيل تقدم موارد حوسبية وخدمات عبر الشبكة. وتحتاج إدارته إلى مراقبة المعالج والذاكرة والتخزين والخدمات والسجلات، لأن المشكلة قد تظهر في مورد واحد بينما يكون سببها الحقيقي خدمة متوقفة أو امتلاء مساحة أو خطأ في الاتصال. لذلك لا يكتفي المشروع بعرض قيمة منفردة، بل يحفظ مجموعة نتائج مرتبطة بالخادم والزمن والأمر المنفذ.")
    add_heading(doc, "2.3 نظام Linux والمراقبة التشغيلية", 2)
    add_text(doc, "تستخدم خوادم المشروع نظام Linux، ولذلك تعتمد المراقبة على أوامر قراءة معروفة مثل معلومات النظام والخدمات والسجلات. لا يسمح التصميم بإرسال shell حر من المستخدم أو النموذج، بل يعرّف الأوامر والمهل وحجم المخرجات مسبقاً. هذا الأسلوب يجعل نتيجة المراقبة قابلة للتكرار والمقارنة ويمنع تحويل وظيفة القراءة إلى قناة تنفيذ غير مقيدة.")
    add_heading(doc, "2.4 الذكاء الصنعي والوكلاء البرمجيون", 2)
    add_text(doc, "الوكيل البرمجي هو برنامج يستقبل هدفاً، يجمع معلومات من البيئة، يطبق سياسة، ثم يعيد نتيجة أو يطلب خطوة تالية. أما النموذج اللغوي فيساعد في تفسير النصوص وربط الملاحظات، لكنه لا يمثل بديلاً عن قواعد النظام. في هذا المشروع يفصل التصميم بين قدرة النموذج على اقتراح الخطوة وبين قدرة Python على التحقق والتنفيذ والتسجيل.")
    add_heading(doc, "2.5 النماذج اللغوية ودور Ollama", 2)
    add_text(doc, "يستخدم Ollama لتحليل التقارير النصية وإنشاء تفسير مساعد للمشغل. ويُعامل الناتج بوصفه نتيجة تحليل قابلة للمراجعة، لا بوصفه أمراً موثوقاً بذاته. فإذا تعطل Ollama أو كانت الاستجابة غير صالحة، يحفظ النظام التقرير وسبب النقص، ولا ينتقل تلقائياً إلى المعالجة.")
    add_heading(doc, "2.6 الاسترجاع المعزز بالسياق RAG", 2)
    add_text(doc, "يعتمد التحليل أحياناً على تقرير حديث أو وثيقة معرفة أو نتيجة فحص متخصص. يربط RAG النتيجة بمصدرها بدلاً من تقديم إجابة معزولة عن السياق. وتستخدم PostgreSQL وpgvector لفهرسة المقاطع واسترجاعها، بينما يحافظ النظام على هوية المصدر وعلاقته بالتقرير والحادثة حتى يستطيع المراجع معرفة سبب ظهور المعلومة.")
    add_heading(doc, "2.7 الأدلة والتحقيق المتخصص", 2)
    add_text(doc, "الدليل في أنظمة الإدارة ليس نصاً عاماً، بل ملاحظة تشغيلية لها مصدر وخادم وزمن ومالك. لذلك تنفذ SpecialistInvestigationLoop تقييم السياسة ثم تستدعي EvidenceCollectionService.collect عبر أدوات القراءة المسموحة. لا يعيد المشروع بناء نظام أدلة جديد؛ بل يحفظ نتيجة الحلقة داخل مخزن التحقيق المستمر ويربطها بالتشخيص النهائي.")
    add_heading(doc, "2.8 بروتوكول MCP وحدود الأدوات", 2)
    add_text(doc, "يقدم Model Context Protocol طريقة منظمة لعرض القدرات إلى Claude Code. في المشروع لا يمثل MCP طرفية عامة، بل قائمة أدوات typed ذات أسماء ومدخلات ونتائج محددة. تمر كل مطالبة عبر التحقق من النوع والسياق والسياسة وسجل التدقيق، ولذلك لا يستطيع النموذج إضافة أداة أو تغيير صلاحية من خلال نص الاستجابة.")
    add_heading(doc, "2.9 الاتصال الآمن عبر SSH", 2)
    add_text(doc, "يستخدم المشروع SSH للوصول إلى الخادم، لكن هوية الهدف تتحقق بواسطة ملف known_hosts. ويقيد الاتصال بأوامر مسجلة ومعاملات محددة ومهلة وحجم مخرجات. لا يكفي نجاح الاتصال لإثبات نجاح المعالجة؛ إذ يلزم فحص مستقل بعد التنفيذ للتأكد من عودة الخادم إلى الحالة المتوقعة.")
    add_heading(doc, "2.10 قاعدة البيانات PostgreSQL وpgvector", 2)
    add_text(doc, "تخزن PostgreSQL الكيانات التشغيلية مثل الخوادم والتقارير والتحقيقات وخطط المعالجة والتفويضات والحجوزات وسجل التدقيق. ويستخدم pgvector للبحث الدلالي في الوثائق والتقارير. ويعتمد التصميم على العلاقات والمفاتيح والفهارس لضمان ملكية النتيجة ومنع ربط دليل بحادثة أخرى.")
    add_heading(doc, "2.11 الأمن والتحكم بالوصول", 2)
    add_text(doc, "تحتاج لوحة الإدارة إلى جلسات محمية وأدوار واضحة وحماية CSRF. يطبق الخادم RBAC على كل طلب، ولا يعتمد على إخفاء زر في الواجهة. كما يمنع النظام تخزين الأسرار في التقرير أو إرسال owner token ومفاتيح SSH إلى المتصفح، وتبقى العمليات الحساسة قابلة للتتبع في سجل التدقيق.")
    add_heading(doc, "2.12 مقارنة بين أساليب الإدارة", 2)
    add_text(doc, "توضح المقارنة التالية الفرق بين إدارة يدوية، ومراقبة آلية تقليدية، ووكيل لغوي غير مقيد، والنموذج المقيد الذي يطبقه المشروع. الغرض من المقارنة تحديد سبب الفصل بين التحليل والتنفيذ، لا الادعاء بأن المشروع يستبدل جميع أدوات الإدارة.")
    add_table(doc, ["الأسلوب", "طريقة العمل", "المشكلة أو القيد", "المعالجة في المشروع"], [
        ["الإدارة اليدوية", "المشغل يقرأ السجلات وينفذ الأوامر بنفسه", "بطيئة وصعبة التتبع وقد تعتمد على الذاكرة", "حفظ التقارير والأدلة وسجل التدقيق"],
        ["مراقبة تقليدية", "أوامر دورية وقواعد ثابتة", "تكتشف الحالة ولا تفسر السياق دائماً", "إضافة تحليل ومصادر وتحقيق عند الحاجة"],
        ["وكيل لغوي غير مقيد", "النموذج يختار أوامر وينفذها مباشرة", "خطر shell حر وغياب ملكية واضحة", "MCP typed وصلاحية تنفيذ داخل Python فقط"],
        ["الوكيل المقيد في المشروع", "مراقبة ثم تحليل ثم دليل ثم سياسة وتنفيذ موثق", "يحتاج بنية تشغيلية وقبولاً حقيقياً", "فشل مغلق واختبارات وقبول منفصل"],
    ], [1900, 2600, 2500, 2360], "الجدول 2: مقارنة أساليب إدارة الخوادم")
    add_heading(doc, "2.13 التقنيات المختارة وسبب الاختيار", 2)
    add_table(doc, ["التقنية", "الدور في المشروع", "سبب الاختيار"], [
        ["Python وFastAPI", "الخدمات وواجهات API", "ملائمة للعقود والاختبارات وتركيب الخدمات"],
        ["PostgreSQL وSQLAlchemy", "البيانات والعلاقات", "ملكية واضحة واستعلامات ومعاملات قابلة للتحقق"],
        ["pgvector", "البحث في السياق", "ربط التقارير والوثائق بالبحث الدلالي"],
        ["Ollama", "تحليل النصوص", "تشغيل محلي مع بقاء القرار خارج النموذج"],
        ["Claude Code وMCP", "الإشراف وعرض القدرات", "تدفق منظم دون منح النموذج سلطة تنفيذ مباشرة"],
        ["SSH وknown_hosts", "الاتصال بالخادم", "تحقق من هوية الهدف وأوامر مسجلة"],
        ["pytest وTestClient", "الاختبار", "اختبارات حتمية للسياسات والواجهات والأمن"],
    ], [2200, 3300, 3860], "الجدول 3: التقنيات وأسباب الاختيار")
    add_heading(doc, "2.14 خلاصة الدراسة النظرية", 2)
    add_text(doc, "تبين الدراسة أن قيمة المشروع لا تأتي من استخدام نموذج لغوي وحده، بل من جمع المراقبة الموثقة والتحليل المرتبط بالمصدر والأدلة والسياسة والحفظ الآمن. ولذلك ينتقل التقرير في الفصل التالي إلى تحديد الممثلين وحالات الاستخدام والمتطلبات التي يجب أن يحققها النظام.")


def chapter2(doc):
    add_heading(doc, "الفصل الثالث: الدراسة التحليلية", 1)
    add_heading(doc, "3.1 الممثلون وحالات الاستخدام", 2)
    add_text(doc, "المستخدمون الأساسيون هم المشرف ومدير النظام والمطوّر. وتتعاون معهم خدمات Python وClaude Code وOllama وقاعدة PostgreSQL وخادم VPS وبيئة الاختبار المعزولة. وتوجد حالات الاستخدام من UC-001 إلى UC-026 في docs/use-cases/use-cases.md.")
    add_heading(doc, "3.1.1 الوصف النصي لحالات الاستخدام", 2)
    add_text(doc, "يوضح الجدول التالي عينة من حالات الاستخدام. يشمل الوصف الممثل والخطوة الأساسية والنتيجة المتوقعة، وهو الأسلوب المستخدم في التقارير الأكاديمية لتوضيح الوظائف بدلاً من الاكتفاء بأسماء المكونات.")
    add_table(doc, ["المعرف", "الممثل", "الوصف المختصر", "النتيجة"], [
        ["UC-001", "المشغل", "يختار الخادم ويشغل المراقبة", "تقرير محفوظ"],
        ["UC-005", "المشغل", "يطلب تحليل تقرير سابق", "تحليل مربوط بالمصدر"],
        ["UC-009", "موجّه التحقيق", "يختار المتخصص وأدوات القراءة", "أدلة وتشخيص"],
        ["UC-013", "خدمة المعالجة", "تبني خطة مرتبطة بالأدلة", "خطة محدودة"],
        ["UC-017", "المشغل والمدير", "يراجعان الخطة ويسجلان الموافقة", "موافقة أو رفض"],
        ["UC-019", "عامل التنفيذ", "يفحص الشروط وينفذ مرة واحدة", "نتيجة وسجل تدقيق"],
    ], [1200, 2200, 3900, 2060], "الجدول 6: عينة من حالات الاستخدام ووصفها")
    add_heading(doc, "3.1.2 مخطط حالات الاستخدام", 2)
    add_text(doc, "يبين المخطط التالي علاقة المشغل والمدير بالوظائف الأساسية. ولا يتصل المستخدم مباشرة بقاعدة البيانات أو بالخادم، بل تمر الطلبات عبر واجهة النظام والقيود المحددة.")
    add_figure(doc, 1, "مخطط حالات الاستخدام", "16-use-case.png")
    add_text(doc, "ولأن المخطط العام قد يخفي بعض التفاصيل، يوضح الشكلان التاليان الوظائف من وجهة نظر كل ممثل على حدة.")
    add_figure(doc, 19, "حالات استخدام المشغل", "19-operator-use-case.png")
    add_figure(doc, 20, "حالات استخدام مدير النظام", "20-admin-use-case.png")
    add_heading(doc, "3.1.3 الوصف التفصيلي للحالات الرئيسية", 2)
    add_text(doc, "تستخدم الحالات التالية لتوضيح العلاقة بين المدخلات والمعالجة والنتيجة. يبدأ كل طلب من ممثل معروف، ويمر عبر خدمة محددة، وينتهي بنتيجة محفوظة أو برفض مفسر.")
    add_table(doc, ["المعرف", "الهدف", "الشروط السابقة", "المسار والنتيجة"], [
        ["UC-001", "تشغيل مراقبة خادم", "خادم وملف مراقبة فعال", "تنفيذ أوامر القراءة وحفظ تقرير بنتيجة كل أمر"],
        ["UC-005", "تحليل تقرير سابق", "تقرير محفوظ ومصدر معروف", "استرجاع السياق أو استدعاء Ollama وحفظ التحليل"],
        ["UC-009", "طلب تحقيق متخصص", "مشكلة تحتاج إلى دليل إضافي", "اختيار متخصص وتقييم السياسة وجمع Evidence"],
        ["UC-013", "إنشاء خطة معالجة", "تشخيص وأدلة كافية", "تحديد الهدف والإجراء والخطر وإنشاء بصمة الخطة"],
        ["UC-017", "الموافقة البشرية", "خطة غير قديمة ومستخدم مخول", "مراجعة الخطة ثم قبولها أو رفضها وتسجيل القرار"],
        ["UC-019", "تنفيذ ذاتي آمن", "سياسة مسموحة واختبار ناجح", "استهلاك تصريح وحجز قصير وتنفيذ ثم تحقق وتدقيق"],
        ["UC-022", "استعادة حجز", "عامل سابق منتهية ملكيته", "فحص السجل ثم استعادة مشروطة أو توقف آمن"],
        ["UC-025", "إدارة لوحة النظام", "جلسة ودور صالحان", "عرض الحالات والقرارات دون كشف الأسرار"],
    ], [1200, 2200, 2700, 3260], "الجدول 7: الوصف التفصيلي لحالات الاستخدام الرئيسية")
    add_heading(doc, "3.1.4 بطاقات حالات الاستخدام", 2)
    add_text(doc, "تقدم البطاقات التالية وصفاً نصياً أقرب إلى أسلوب التقارير السابقة. يبين كل وصف الممثلين والشروط والمسار الطبيعي والنتيجة والمسارات البديلة، حتى يمكن الانتقال من اسم الحالة في المخطط إلى سلوك قابل للفحص.")
    use_case_cards = [
        ("UC-001: تشغيل مراقبة خادم", "المشغل", "اختيار خادم وملف مراقبة فعال، وتوفر جلسة صالحة.", "تقرير مراقبة محفوظ مرتبط بالخادم ووقت التنفيذ.", [("المشغل", "يفتح صفحة المراقبة ويختار الخادم والملف."), ("النظام", "يتحقق من الدور ووجود الملف والأوامر المسجلة."), ("النظام", "ينفذ أوامر القراءة ضمن المهلة وحجم المخرجات."), ("النظام", "يحفظ نتيجة كل أمر ثم ينشئ التقرير ويعرض حالته.")], ["E1: إذا فشل الاتصال، تحفظ النتيجة كفشل ولا ينشأ تقرير ناجح.", "E2: إذا انتهت الجلسة، يرفض الطلب قبل تنفيذ أي أمر."]),
        ("UC-005: تحليل تقرير محفوظ", "المشغل وAnalysisOrchestrator", "وجود تقرير محفوظ ومصدر معروف وعدم وجود تحليل صالح قابل لإعادة الاستخدام.", "تحليل محفوظ مع المراجع ومؤشر اكتمال النتيجة.", [("المشغل", "يختار التقرير ويطلب التحليل."), ("النظام", "يفحص التقرير والتحليل السابق ومصادر السياق."), ("النظام", "يعيد استخدام نتيجة مناسبة أو يرسل سياقاً محدوداً إلى Ollama."), ("النظام", "يحفظ التحليل ويربطه بالتقرير ويعرض النتيجة.")], ["E1: عند تعطل Ollama يحفظ النظام سبب الفشل ويبقي التقرير متاحاً.", "E2: عند نقص المصدر يعرض التحليل كغير مكتمل ولا يقترح معالجة."]),
        ("UC-009: طلب تحقيق متخصص", "المشغل وInvestigationRouter", "وجود حادثة تحتاج إلى دليل إضافي ووجود متخصّص فعال بأدوات مسجلة.", "نتيجة تحقيق مستمرة تحتوي Evidence ومصادرها واختلافات المتخصصين.", [("المشغل", "يطلب فحصاً إضافياً من صفحة الحادثة."), ("الموجّه", "يختار المتخصصين وفق المشكلة والسياسة والميزانية."), ("حلقة التحقيق", "تشغل أدوات القراءة المحددة وتستدعي EvidenceCollectionService.collect."), ("النظام", "يحفظ الأدلة ويربطها بالحادثة والخادم ويصدر نتيجة التحقيق.")], ["E1: يرفض النظام أداة غير مسجلة أو دليلاً يخص خادماً آخر.", "E2: عند اختلاف النتائج، يعرض الاختلاف ويطلب مراجعة بشرية."]),
        ("UC-013: إنشاء خطة معالجة", "خدمة المعالجة والمشغل", "وجود تشخيص وأدلة كافية وإجراء معروف وهدف محدد.", "خطة ثابتة تحمل بصمة ومخاطر وهدفاً وشروط تحقق.", [("النظام", "يقرأ التشخيص والأدلة ويحدد الإجراء المقترح."), ("النظام", "يتحقق من الهدف والخادم والسياسة ومستوى الخطورة."), ("النظام", "ينشئ الخطة وبصمتها وشروط ما قبل التنفيذ."), ("المشغل", "يراجع الخطة قبل نقلها إلى الموافقة أو الاختبار.")], ["E1: إذا تغيرت الأدلة أو الهدف، تصبح الخطة قديمة.", "E2: عند نقص دليل حاسم، تتوقف عملية إنشاء الخطة."]),
        ("UC-017: الموافقة البشرية", "المشغل ومدير النظام", "خطة غير قديمة ومستخدم يملك صلاحية الموافقة وحماية CSRF صالحة.", "قرار موافقة أو رفض مرتبط بالخطة الحالية وسجل تدقيق.", [("المشغل", "يفتح تفاصيل الخطة ويقرأ الخطر والهدف والدليل."), ("النظام", "يتحقق من الجلسة والدور وبصمة الخطة."), ("المشغل", "يختار الموافقة أو الرفض ويدخل السبب عند الحاجة."), ("النظام", "يحفظ القرار ويمنع استخدامه مع خطة مختلفة.")], ["E1: يرفض النظام القرار إذا تغيرت الخطة بعد فتح الصفحة.", "E2: يرفض طلباً دون CSRF أو بدور غير كافٍ."]),
        ("UC-019: تنفيذ ذاتي آمن", "عامل التنفيذ وPolicyEvaluator", "سياسة مسموحة، تصريح غير مستهلك، خطة مطابقة، اختبار ناجح، وحجز متاح.", "تنفيذ واحد بنتيجة تحقق وسجل تدقيق أو رفض مفسر.", [("النظام", "يطابق السياسة والبصمات ونتيجة الاختبار وسجل المحاولات."), ("النظام", "يستهلك التصريح وينشئ حجزاً قصير المدة وowner token."), ("العامل", "ينفذ العملية المسماة خارج معاملة قاعدة البيانات."), ("النظام", "يتحقق من النتيجة ثم يدمجها ذرّياً إذا بقي العامل مالكاً.")], ["E1: تمنع البصمة أو التصريح المستهلك إعادة التنفيذ.", "E2: عند انتهاء الحجز أو تغير المالك، يرفض الحفظ المتأخر."]),
        ("UC-022: استعادة حجز منتهٍ", "عامل التعافي ومدير النظام", "حجز منتهٍ وعدم وجود تنفيذ نهائي واضح أو مالك نشط.", "استعادة مشروطة لحجز آمن أو إيقاف العملية لطلب مراجعة.", [("عامل التعافي", "يقرأ حالة الحجز والتصريح وسجل التنفيذ."), ("النظام", "يفحص مدة الانتهاء ووجود مالك آخر ووضوح النتيجة."), ("النظام", "يسترد الحجز وفق الشرط أو يضع الحالة في انتظار المراجعة."), ("مدير النظام", "يراجع الحالة الغامضة ولا يعيد أمراً غير واضح.")], ["E1: إذا كان التنفيذ قد بدأ ولا يمكن معرفة نتيجته، يتوقف النظام.", "E2: لا يسمح باسترداد حجز يملكه عامل نشط."]),
        ("UC-025: إدارة لوحة النظام", "مدير النظام", "جلسة server-side صالحة ودور يسمح بالصفحة المطلوبة.", "عرض بيانات تشغيلية مصفاة مع سجل تدقيق لكل عملية حساسة.", [("المدير", "يسجل الدخول ويفتح لوحة الحالة أو السياسة أو التدقيق."), ("النظام", "يتحقق من الجلسة والدور وCSRF عند الطلبات المعدلة."), ("النظام", "يقرأ الإسقاطات المسموحة من الخدمات ولا يعرض الأسرار."), ("المدير", "يراجع الحالة أو يغير سياسة مسموحة ويسجل السبب.")], ["E1: يعاد المستخدم إلى صفحة الدخول عند انتهاء الجلسة.", "E2: يحجب النظام الحقول الحساسة حتى لو طلبها العميل."]),
    ]
    for index, (title, actor, preconditions, postconditions, steps, alternatives) in enumerate(use_case_cards, 1):
        add_heading(doc, f"حالة الاستخدام {index}: {title}", 3)
        add_text(doc, f"الممثلون: {actor}")
        add_text(doc, f"الشروط السابقة: {preconditions}")
        add_text(doc, "المسار الأساسي:")
        add_table(doc, ["الممثل", "الاستجابة أو المعالجة"], steps, [2500, 7100])
        add_text(doc, f"الشروط اللاحقة: {postconditions}")
        add_text(doc, "المسارات البديلة والاستثناءات:")
        for alternative in alternatives:
            add_bullet(doc, alternative)
    add_heading(doc, "3.2 دورة الحادثة", 2)
    add_text(doc, "تبدأ دورة العمل بملاحظة أو مشكلة في الخادم. يجمع النظام تقريراً، ثم يحلله، ويطلب فحصاً إضافياً إذا لم تكن المعلومات كافية. بعد ذلك تظهر نتيجة التشخيص والأدلة، ويمكن إنشاء خطة معالجة. إما أن يوافق عليها المشرف، أو تمر في مسار التنفيذ الذاتي إذا تحققت شروطه. وفي حال الفشل يتوقف النظام ويسجل السبب.")
    add_heading(doc, "3.3 نموذج الخطورة", 2)
    add_table(doc, ["القرار", "المعنى", "الشرط التقريبي"], [["تنفيذ ذاتي", "تنفيذ محدود", "نجاح السياسة والاختبار والتاريخ والأدلة وحدود المحاولات"], ["موافقة بشرية", "لا تنفيذ ذاتي", "الحالة تحتاج إلى موافقة أو الثقة غير كافية"], ["رفض", "إيقاف العملية", "خطر مرتفع أو بيانات غير متطابقة أو اختبار مفقود أو محاولة تكرار"]], [2200, 2600, 4360], "الجدول 8: قرارات سياسة المعالجة الذاتية")
    add_heading(doc, "3.4 التحكم البشري", 2)
    add_text(doc, "الموافقة تسجل في قاعدة البيانات وترتبط بالخطة الحالية. وقبل التنفيذ يفحص النظام هوية المستخدم وصلاحيته وحماية CSRF. يستطيع مدير النظام إيقاف السياسة أو تشغيلها، لكنه لا يستطيع من الواجهة تجاوز الفحص أو إنشاء تصريح يدوي.")
    add_heading(doc, "3.5 وصف تدفقات النظام", 2)
    add_text(doc, "في المراقبة يختار النظام ملفاً محفوظاً، وينفذ أوامر القراءة المسجلة، ثم يحفظ النتائج في تقرير. بعد ذلك يستخدم التقرير في التحليل، أو يعيد استخدام تحليل سابق إذا كان مناسباً.")
    add_text(doc, "في التحقيق يختار النظام المتخصص المناسب للمشكلة. يعمل المتخصص بأدوات القراءة المسموحة، ثم تجمع الحلقة الأدلة وتربطها بالحادثة والخادم. ولا يستطيع المتخصص تغيير السياسة أو تشغيل أمر غير مسجل.")
    add_text(doc, "في المعالجة الذاتية يفحص النظام السياسة ونتيجة الاختبار وسجل المحاولات وتفاصيل الخطة. ثم يستخدم تصريحاً لمرة واحدة ويحجز العملية مدة قصيرة. ينفذ الإجراء خارج معاملة قاعدة البيانات، وبعد ذلك يتحقق من النتيجة ويسجلها إذا بقي العامل مالكاً للعملية.")
    add_heading(doc, "3.6 السيناريو التشغيلي الطبيعي", 2)
    add_text(doc, "يفتح المشرف لوحة الإدارة، ويختار الخادم وملف المراقبة، ثم يراجع التقرير. إذا ظهرت مشكلة، يطلب النظام تحليلاً وفحصاً إضافياً. وتعرض الواجهة التشخيص والأدلة والمصادر حتى يستطيع المشرف فهم النتيجة.")
    add_text(doc, "إذا احتاجت المشكلة إلى معالجة، ينشئ النظام خطة واضحة. يراجعها المشرف ويوافق عليها في المسار العادي. أما المسار الذاتي فلا ينفذ إلا بعد نجاح جميع الفحوص، وتبقى النتيجة ظاهرة وقابلة للمراجعة.")
    add_heading(doc, "3.7 سيناريوهات الفشل والاستجابة", 2)
    add_table(doc, ["الحالة", "الاستجابة المتوقعة", "السبب"], [["تعطل PostgreSQL", "إيقاف الطلب وتسجيل الفشل", "لا قرار دون قاعدة البيانات"], ["تعطل Ollama", "حفظ التقرير وإظهار نتيجة غير مكتملة", "النموذج لا ينفذ الأوامر"], ["فقدان known_hosts", "رفض اتصال SSH", "منع اتصال غير موثق"], ["اختلاف الأدلة", "عرض الاختلاف وطلب المراجعة", "منع التشخيص الخاطئ"], ["انتهاء الحجز", "فحص السجل واسترداد الحجز عند السماح", "منع الحفظ المتأخر"], ["إعادة استخدام التصريح", "رفض الطلب وتسجيل الحدث", "التصريح يستخدم مرة واحدة"], ["تكرر الفشل", "إيقاف المحاولات مؤقتاً", "تقليل الضرر"]], [2600, 3900, 3300], "الجدول 9: تحليل سيناريوهات الفشل")
    add_heading(doc, "3.8 المتطلبات الأمنية في التحليل", 2)
    add_text(doc, "يستخدم النظام النموذج اللغوي للمساعدة في الفهم فقط. أما القرار والتنفيذ فهما من مسؤولية Python وقاعدة البيانات. لذلك لا يستطيع النص الناتج من Claude أو Ollama تغيير الصلاحيات أو إصدار تصريح أو إضافة أمر جديد.")
    add_text(doc, "يفترض التصميم أن يعيد المستخدم إرسال الطلب، أو أن تكون الخطة قديمة، أو أن يتغير الخادم والهدف، أو أن يحاول عامل قديم حفظ نتيجة بعد انتهاء ملكيته. وتمنع هذه الحالات البصمة والحجز وCSRF والصلاحيات وسجل التدقيق. وإذا لم ينجح أي فحص، يرفض النظام العملية.")
    add_heading(doc, "3.9 معايير قابلية الاستخدام والمراجعة", 2)
    add_text(doc, "يجب أن يستطيع المشرف فهم ما حدث دون قراءة الشيفرة. لذلك تعرض لوحة الإدارة حالة العملية والقرار والنتيجة والوقت وسبب الرفض أو التأجيل. ولا تعرض كلمات المرور أو الرموز السرية أو مفاتيح الملكية، كما تفصل بين صلاحية المشاهدة وصلاحية التشغيل.")
    add_heading(doc, "3.10 قرارات التصميم التحليلية", 2)
    add_text(doc, "أهم قرار في التصميم هو فصل الاقتراح عن التنفيذ. كما تحفظ الأدلة والتشخيص في قاعدة البيانات بدلاً من الاعتماد على محادثة مؤقتة. وفصل الحجز عن التنفيذ يجعل الاتصال الخارجي قصيراً وآمناً، ويمنع تعارض عاملين في الوقت نفسه.")


def chapter3(doc):
    add_heading(doc, "الفصل الرابع: الدراسة التصميمية", 1)
    add_heading(doc, "4.1 مخطط النظام وسياقه", 2)
    add_figure(doc, 4, "مخطط النظام العام", "01-system-block.png")
    add_figure(doc, 5, "سياق النظام", "02-system-context.png")
    add_heading(doc, "4.2 المخططات الزمنية ومخطط النشاط", 2)
    add_text(doc, "توضح المخططات التالية طريقة انتقال الطلب بين المستخدم والواجهات وخدمات المراقبة والتحليل وقاعدة البيانات. ويبين مخطط النشاط مسار المعالجة الخاضعة للإشراف من قراءة الأدلة حتى التحقق أو الاستعادة عند الفشل.")
    add_figure(doc, 2, "مخطط تسلسل المراقبة والتحليل", "17-monitoring-sequence.png")
    add_figure(doc, 21, "مخطط تسلسل التحقيق المتخصص", "21-specialist-sequence.png")
    add_figure(doc, 22, "مخطط تسلسل المعالجة الخاضعة للإشراف", "22-supervised-sequence.png")
    add_figure(doc, 23, "مخطط تسلسل التنفيذ الذاتي", "23-autonomous-sequence.png")
    add_figure(doc, 24, "مخطط تسلسل موافقة المدير والتدقيق", "24-admin-sequence.png")
    add_figure(doc, 3, "مخطط نشاط المعالجة الخاضعة للإشراف", "18-remediation-activity.png")
    add_heading(doc, "4.3 طبقات المكونات", 2)
    add_figure(doc, 6, "طبقات المكونات والاعتماديات", "03-component-layers.png")
    add_text(doc, "يتكون النظام من واجهات للمستخدم، وخدمات تنفذ وظائف المشروع، وطبقة للبنية التحتية. توجد العقود والسياسات المشتركة في app/core. وتربط طبقة composition هذه الأجزاء معاً، ولا تنشئ مساراً آخر لتنفيذ الأوامر.")
    add_heading(doc, "4.4 Claude Code وOllama وMCP", 2)
    add_figure(doc, 7, "تدفق Claude Code -> Ollama -> MCP -> Python", "04-claude-ollama-mcp-python.png")
    add_text(doc, "يحدد Claude الخطوة المطلوبة، ثم يرسل الطلب عبر MCP إلى Python. تفحص Python المدخلات والصلاحيات قبل استدعاء Ollama للتحليل أو خدمة البنية التحتية. وفي النهاية تعيد نتيجة منظمة إلى الواجهة.")
    add_heading(doc, "4.5 المراقبة والتحليل والتحقيق", 2)
    add_figure(doc, 8, "المراقبة والتحليل والتحقيق", "05-monitor-analysis-investigation.png")
    add_heading(doc, "4.6 المتخصصون وEvidence", 2)
    add_figure(doc, 9, "تنسيق SpecialistInvestigationLoop", "06-specialist-orchestration.png")
    add_text(doc, "تنفذ حلقة التحقيق فحص السياسة ثم تجمع الأدلة من خلال EvidenceCollectionService.collect. وتحفظ قاعدة البيانات النتائج ومصادرها وروابطها. لذلك لا نحتاج إلى بناء نظام أدلة جديد، بل نربط نتيجة الحلقة بمخزن التحقيق المستمر.")
    add_heading(doc, "4.7 المعالجة الخاضعة للإشراف وSandbox", 2)
    add_figure(doc, 10, "سلسلة المعالجة الخاضعة للإشراف", "07-supervised-remediation.png")
    add_figure(doc, 11, "سلسلة Sandbox والتحقق والاستعادة", "08-sandbox-validation.png")
    add_heading(doc, "4.8 المعالجة الذاتية والحجز", 2)
    add_figure(doc, 12, "سلسلة المعالجة الذاتية", "09-autonomous-sequence.png")
    add_figure(doc, 13, "السياسة والتفويض والحجز وعدم التكرار", "10-policy-reservation-idempotency.png")
    add_text(doc, "الحجز مؤقت ولا يبقى مفتوحاً أثناء الاتصال بـ SSH أو Ollama. ويضمن owner_token أن العامل القديم لا يحفظ نتيجة عامل جديد. كما يمنع المفتاح الفريد للعملية تنفيذها مرتين.")
    add_heading(doc, "4.9 قاطع الدارة والاستعادة", 2)
    add_figure(doc, 14, "قاطع الدارة والاستعادة", "11-circuit-recovery.png")
    add_heading(doc, "4.10 Admin وRBAC والتدقيق", 2)
    add_figure(doc, 15, "مصادقة Admin وRBAC وCSRF", "12-admin-auth-rbac.png")
    add_text(doc, "تستخدم خدمة تسجيل الدخول كلمات مرور محمية، وجلسات محفوظة في الخادم، وحماية CSRF، وثلاثة أدوار هي المشاهد والمشغل والمدير. وتسجل الأحداث المهمة، بينما تخفي الشاشة رموز التفويض ومفتاح ملكية العملية.")
    add_heading(doc, "4.11 النشر وقاعدة البيانات والمكونات", 2)
    add_figure(doc, 16, "البنية الفيزيائية", "13-deployment.png")
    add_figure(doc, 17, "مخطط علاقات قاعدة البيانات", "14-database-erd.png")
    add_figure(doc, 18, "المكونات والفئات الأساسية", "15-key-components.png")
    add_heading(doc, "4.12 تصميم العقود والحدود", 2)
    add_text(doc, "تحتوي app/core على العقود التي تصف الحالات والنتائج والسياسات، بينما تطبق طبقة capabilities السلوك التشغيلي. هذا الترتيب يمنع أن تتسرب تفاصيل SQL أو SSH أو عميل النموذج إلى الواجهات. كما يسمح للاختبارات باستخدام كائنات بديلة مع بقاء شكل القرار والنتيجة ثابتاً.")
    add_text(doc, "حد MCP ليس مجرد قناة نقل؛ إنه قائمة قدرات typed ذات أسماء ومدخلات ونتائج محددة. يستطيع Claude طلب عملية معروفة، لكن Python تتحقق من المدخلات والسياق وتعيد نتيجة منظمة. أي خدمة جديدة يجب أن تمر عبر هذا الحد وأن تلتزم بسياسة التسجيل، بدلاً من إضافة طرفية عامة أو حقل نصي يحمل أوامر.")
    add_heading(doc, "4.13 تصميم طبقة التخزين", 2)
    add_text(doc, "تتبع النماذج قاعدة ملكية واضحة: الخادم يملك سياق المراقبة، التقرير يملك التحليل، التحقيق يملك الأدلة والتشخيص، والخطة تملك المعالجة والتحقق. الروابط لا تستخدم للتجميل، بل تمنع خلط نتيجة من خادم أو حادثة مع طلب آخر. وتوفر المفاتيح والفهارس القيود التي يصعب ضمانها في الذاكرة وحدها.")
    add_text(doc, "تُستخدم PostgreSQL للبيانات التشغيلية وللبحث الدلالي عبر pgvector، مع فهارس مخصصة للنطاق والمصدر والاسترجاع. أما migrations فتبقى إضافية لتجنب فقد البيانات، ويظل bootstrap قادراً على التحقق من الجداول المتوقعة. توثيق المجموعات في التقرير لا يستبدل فحص المخطط؛ بل يشرح كيف ترتبط الجداول بالوظائف.")
    add_heading(doc, "4.14 تصميم الاتساق والتزامن", 2)
    add_text(doc, "تتعامل المنصة مع التزامن على أساس أن عاملين قد يقرآن الحالة نفسها في اللحظة نفسها. ولذلك لا تعتمد على قراءة snapshot ثم كتابة metadata كاملة دون شرط. الحجز والتحديث المشروط والمالك المؤقت تجعل كل عامل يثبت فقط النتيجة التي يملك حق تثبيتها، بينما تعيد العملية المتنافسة قراءة الحالة أو ترفضها.")
    add_text(doc, "هذه القاعدة مهمة خصوصاً عند تجميع Evidence أو تحديث حالة معالجة ذاتية. يجب أن يحافظ الدمج على النتائج السابقة، وألا يستبدل قاموساً كاملاً نتيجة قراءة قديمة. أما الحقول التي تتطلب تغييراً وحيداً، مثل الحالة النهائية أو استهلاك التصريح، فتُحدّث داخل عملية ذرية مناسبة وتُراجع نتيجتها.")
    add_heading(doc, "4.15 تصميم الأمن والخصوصية", 2)
    add_text(doc, "يعتمد Admin على جلسات server-side وdigest للكوكي وكلمات مرور مع scrypt وCSRF وRBAC. وتخضع نقاط API وWeb إلى middleware مشترك حتى لا تصبح إخفاءات الواجهة بديلاً عن الحماية. وتستخدم إسقاطات العرض قائمة حقول صريحة، ولذلك لا تصل رموز التفويض أو owner token أو المفاتيح الخاصة إلى المتصفح.")
    add_text(doc, "في طبقة SSH، known_hosts جزء من هوية الهدف وليس إعداداً اختيارياً في مسار التشغيل الآمن. وتُحصر الأوامر في سجل معروف وتُمرر المعاملات وفق عقد القدرة، مع منع shell الحر. كما تسجل الأحداث الأمنية ومحاولات الرفض بما يسمح بتحليل الحادثة من دون تخزين أسرار في التقرير.")
    add_heading(doc, "4.16 تصميم القابلية للتوسع", 2)
    add_text(doc, "يمكن توسيع عدد الخوادم أو المتخصصين أو السياسات دون إضافة فرع خاص لكل حالة، لأن التعريفات والميزانيات والروابط محفوظة في قاعدة البيانات. ويجب أن يحافظ التوسع على حدود المعدل والـlease وسجل التدقيق، لا أن يلتف عليها. كما يمكن إضافة مزود نموذج أو واجهة جديدة عبر composition مع إبقاء Python صاحبة قرار التحقق.")
    add_heading(doc, "4.17 الأنماط والمبادئ التصميمية", 2)
    add_text(doc, "يعتمد التصميم على فصل المسؤوليات وحقن الاعتماديات وطبقة المستودعات وواجهات العقود. تظهر هذه المبادئ في composition الذي يجمع الخدمات، وفي المستودعات التي تخفي تفاصيل SQL، وفي العقود التي تثبت شكل القرار والنتيجة. وتفيد هذه البنية في استبدال عميل Ollama أو مصدر SSH في الاختبارات دون تغيير سياسة التنفيذ.")
    add_table(doc, ["المبدأ أو النمط", "موضع الاستخدام", "الفائدة"], [
        ["Layered architecture", "core وcapabilities وinfrastructure وinterfaces", "منع تسرب تفاصيل البنية إلى الواجهة"],
        ["Repository", "مستودعات التقارير والسياسات والتنفيذ", "عزل الاستعلامات وتوحيد الحفظ"],
        ["Dependency composition", "app/composition", "تركيب نسخة واحدة من الخدمات والاعتماديات"],
        ["Strategy", "اختيار المتخصص أو مسار التحليل", "تغيير السلوك عبر تعريفات دون فروع متكررة"],
        ["State machine", "السياسة والحجز وقاطع الدارة", "منع الانتقال غير الصالح وتسجيل الحالة"],
        ["Fail-closed", "التنفيذ والسياسة والاتصال", "إيقاف المسار عند نقص الدليل أو غموض النتيجة"],
    ], [2200, 3400, 3560], "الجدول 10: المبادئ والأنماط التصميمية")
    add_heading(doc, "4.18 تصميم الواجهات والعقود", 2)
    add_text(doc, "تتواصل لوحة Admin وأدوات MCP مع خدمات الخادم عبر عقود typed. يحتوي الطلب على هوية العملية والسياق والمدخلات المسموحة، بينما تحتوي الاستجابة على الحالة والنتيجة وسبب الرفض ومراجع الأدلة عند توفرها. ولا تعتمد الواجهة على شكل داخلي لجدول قاعدة البيانات، بل تحصل على إسقاط مصمم للعرض.")
    add_text(doc, "تستخدم الواجهة نماذج منفصلة للمراقبة والتحليل والسياسة والحجز والتدقيق. وتعرض كل صفحة الحالة الحالية مع وقت آخر تحديث، بينما تعيد API أخطاء منظمة يمكن اختبارها. هذا الفصل يجعل تغيير قالب Admin أو إضافة عميل آخر ممكناً دون نقل السلطة التنفيذية إلى العميل.")
    add_heading(doc, "4.19 الأدوات والبيئات المستخدمة", 2)
    add_table(doc, ["الأداة أو البيئة", "الاستخدام في المشروع", "الحد أو القيد"], [
        ["Python وFastAPI", "الخدمات وواجهات API والعقود", "لا تمنح النموذج سلطة تنفيذ"],
        ["PostgreSQL وSQLAlchemy", "البيانات التشغيلية والمعاملات والمستودعات", "يجب توفر المخطط والاتصال قبل القرار"],
        ["pgvector", "استرجاع مقاطع التقارير والوثائق", "لا يستبدل مصدر الدليل أو التحقق"],
        ["Ollama", "تحليل النصوص وإنشاء التضمينات", "تعطل الخدمة ينتج مساراً غير مكتمل"],
        ["Claude Code وMCP", "الإشراف وعرض الأدوات typed", "لا يوجد shell حر أو اتصال SSH مباشر"],
        ["SSH وknown_hosts", "الوصول الموثق إلى VPS", "لا يقبل الهدف غير المعروف"],
        ["Jinja2 وJavaScript وCSS", "صفحات Admin والقوالب والتفاعلات", "السلطة النهائية في backend"],
        ["pytest وTestClient", "الاختبارات الحتمية والمحلية", "القبول الخارجي opt-in فقط"],
    ], [2300, 4700, 3160], "الجدول 11: الأدوات والبيئات وحدود استخدامها")
    add_heading(doc, "4.20 بيئة التشغيل والنشر", 2)
    add_text(doc, "تعمل بيئة التطوير على Windows وWSL2، بينما يتصل المسار التشغيلي بخادم Linux عبر SSH موثق. تحتاج الخدمة إلى PostgreSQL وملف known_hosts ومفتاح SSH خارج المستودع، كما يحتاج التحليل إلى Ollama عند تشغيله فعلياً. تحفظ الإعدادات في متغيرات البيئة ولا تكتب الأسرار داخل التقرير أو ملفات الرسوم.")
    add_text(doc, "يبدأ النشر بفحص إعدادات قاعدة البيانات والمخطط والاتصال، ثم فحص الأدوات والملفات الخارجية. وتبقى المعالجة الذاتية معطلة افتراضياً حتى تكتمل السياسة والهدف الآمن. وتفصل اختبارات real_runtime عن regression المحلي حتى لا تصل الاختبارات العادية إلى خادم حقيقي.")
    add_heading(doc, "4.21 تصميم واجهات API وMCP", 2)
    add_text(doc, "تقسم الواجهات إلى طلبات مراقبة وتقارير وتحليل وتحقيق ومعالجة وإدارة. يحتوي كل طلب على مدخلات محددة، وتعيد الخدمة حالة منظمة تتضمن النتيجة أو سبب الرفض أو النقص. أما MCP فيعرض القدرات المسجلة فقط، وتتحقق Python من النوع والسياق والمالك قبل استدعاء أي قدرة.")
    add_text(doc, "تستخدم واجهة Admin إسقاطات خاصة للعرض ولا تمرر نماذج قاعدة البيانات كاملة إلى المتصفح. وتخضع الطلبات المعدلة للجلسة والدور وCSRF. بهذا الشكل تبقى الواجهة وسيلة عرض ومراجعة، بينما تبقى الخدمات والقاعدة مصدر القرار والتنفيذ.")


def chapter4(doc):
    add_text(doc, "في إعداد الإنتاج تُفرض قيمة ADMIN_SESSION_SECRET من خارج المستودع وبطول لا يقل عن 32 محرفاً، مع ADMIN_SESSION_SECURE=true. ويكون النشر خلف HTTPS reverse proxy، وتبقى PostgreSQL وOllama داخلية أو خاصة، ويظل حد MCP محدوداً typed، ويُفرض SSH known_hosts، ويكون السماح بالمعالجة الذاتية false افتراضياً.")
    add_heading(doc, "الفصل الخامس: التنفيذ والاختبارات", 1)
    add_heading(doc, "5.1 البيئة والتقنيات وسبب الاختيار", 2)
    add_table(doc, ["التقنية", "الاستخدام والسبب"], [["Python وFastAPI", "تنفيذ خدمات المشروع وبناء واجهات قابلة للاختبار."], ["Jinja2 وJavaScript وCSS", "بناء لوحة إدارة بسيطة."], ["PostgreSQL وpgvector", "حفظ البيانات والأدلة والبحث في الوثائق."], ["SQLAlchemy وpsycopg", "ربط الخدمات بقاعدة البيانات."], ["Ollama", "تحليل التقارير وإنشاء التضمينات."], ["Claude Code وMCP", "الإشراف والوصول إلى الأدوات المسموحة."], ["SSH وknown_hosts", "اتصال موثق بالخادم وبأوامر مسجلة."], ["pytest وTestClient", "اختبار السياسات والواجهات والأمان."]], [2600, 6760], "الجدول 12: التقنيات الرئيسية")
    add_heading(doc, "5.2 تنفيذ قاعدة البيانات", 2)
    add_text(doc, "ينشئ برنامج تهيئة قاعدة البيانات الجداول والفهارس المطلوبة، ويتأكد من وجود pgvector. ثم يقارن الجداول الموجودة بالقائمة المتوقعة. يبلغ عدد الجداول الحالي 33 جدولاً. وتشرح وثيقة بنية البيانات وظيفة كل مجموعة، بينما تضيف ملفات SQL التغييرات دون حذف البيانات.")
    add_heading(doc, "5.3 تنفيذ التكامل", 2)
    add_text(doc, "تربط طبقة composition المستودعات بالخدمات ثم ببيئة التشغيل. وتستخدم لوحة الإدارة وأدوات MCP الخدمات نفسها، لذلك لا يوجد مسار ثانٍ يخالف المسار الأساسي. لا يملك Ollama صلاحية التنفيذ، ولا يتصل Claude بخادم SSH مباشرة.")
    add_heading(doc, "5.4 Admin UI", 2)
    add_text(doc, "تضم لوحة الإدارة صفحات المراقبة والسياسات والحجوزات والتفويضات وسجل التدقيق. وتستخدم طلبات الواجهة طريقة موحدة لإرسال CSRF وقراءة الأخطاء. وتبقى الصلاحية الحقيقية في الخادم، أما إخفاء زر في الواجهة فليس وسيلة حماية.")
    add_heading(doc, "5.5 شرح تنفيذي لدورة المراقبة والتحقيق", 2)
    add_text(doc, "تبدأ المراقبة من ملف محفوظ في قاعدة البيانات. يحدد الملف الخادم والأوامر والمهلة وحجم المخرجات. تنفذ الخدمة أوامر القراءة فقط، وتحفظ نتيجة كل أمر مع وقت التنفيذ وحالته. وبذلك يمكن الرجوع إلى التقرير لاحقاً.")
    add_text(doc, "بعد حفظ التقرير، يبحث النظام عن تحليل سابق مناسب. وإذا لم يجده، يرسل المعلومات المطلوبة إلى Ollama مع مصادرها. يساعد Ollama في تفسير النتائج، لكنه لا يستطيع تنفيذ أي أمر. وتبقى التقارير المحفوظة هي المرجع الأساسي.")
    add_text(doc, "إذا احتاجت المشكلة إلى فحص إضافي، يختار InvestigationRouter المتخصص المناسب. تنفذ حلقة التحقيق أدوات القراءة المسموحة، ثم تجمع Evidence من خلال EvidenceCollectionService.collect وتحفظها مع الحادثة والخادم والمصدر والوقت.")
    add_text(doc, "بعد انتهاء المتخصصين، يقارن النظام النتائج ويظهر النقص أو الاختلاف. ثم ينشئ تشخيصاً يبين المشكلة والأدلة التي اعتمد عليها. ولا تتحول نتيجة المتخصص إلى أمر تنفيذ مباشر.")
    add_heading(doc, "5.6 شرح تنفيذي لدورة المعالجة الذاتية", 2)
    add_text(doc, "المعالجة الذاتية منفصلة عن المراقبة والتحقيق. يفحص النظام السياسة والخادم ونوع الإجراء وبصمة الخطة ونتيجة Sandbox وسجل المحاولات. وإذا كانت معلومة واحدة ناقصة أو غير متطابقة، يرفض التنفيذ.")
    add_text(doc, "قبل التنفيذ يصدر النظام تصريحاً لمرة واحدة. يرتبط التصريح بالقرار والسياسة والخطة والخادم والإجراء ونتيجة الاختبار. ولا يمكن استخدام التصريح مرة ثانية بعد استهلاكه.")
    add_text(doc, "يحجز النظام العملية مدة قصيرة، ثم ينفذ الإجراء خارج معاملة قاعدة البيانات. وهذا مهم لأن الاتصال بـ SSH قد يستغرق وقتاً. بعد التنفيذ يفحص النظام النتيجة ويحفظ التنفيذ والتحقق والتراجع عند الحاجة.")
    add_text(doc, "يرتبط الحفظ النهائي بمالك العملية. فإذا انتهى الحجز أو استرده عامل آخر، لا يستطيع العامل القديم حفظ نتيجته. ويسجل النظام كل خطوة في سجل التدقيق.")
    add_text(doc, "يوقف قاطع الدارة المحاولات عند تكرر الفشل. وعند عودة النظام بعد انقطاع، يفحص الحجز والتصريح وسجل التنفيذ قبل المتابعة. وإذا كانت الحالة غير واضحة، يتوقف ويطلب مراجعة بدلاً من إعادة عملية خطرة.")
    add_heading(doc, "5.7 شرح تنفيذ الاختبارات والتكامل", 2)
    add_text(doc, "تغطي الاختبارات العقود والسياسات والمستودعات والواجهات والأمان. وتفحص اختبارات الوحدات قرارات السياسة وحدود الأوامر وتجميع الأدلة. كما تفحص اختبارات المستودعات التزامن والحجز والتصريح ومنع التكرار.")
    add_text(doc, "يتحقق اختبار قاعدة البيانات من الجداول وpgvector والفهارس المطلوبة. أما اختبار القبول الحقيقي فيحتاج Ollama وSSH وبيانات تشغيل حقيقية، ولذلك يبقى منفصلاً عن الاختبارات المحلية.")
    add_text(doc, "تشمل الخطة اختبارات الترجمة، والسياسات، وقاعدة البيانات، وواجهات API، ولوحة الإدارة، وأدوات MCP، والأمان، والتزامن، والاستعادة، وقاطع الدارة. وقد فصلنا بين نتائج الاختبارات المحلية ونتيجة القبول الحقيقي.")
    add_heading(doc, "5.8 نتائج الاختبار الحالية", 2)
    add_table(doc, ["الفحص", "النتيجة"], [["الاختبارات المحلية", "586 ناجحاً، دون فشل أو تخطٍ، مع تحذير واحد"], ["قاعدة البيانات", "33 من 33، وpgvector موجود، و3 فهارس للبحث"], ["أدوات MCP", "25 أداة"], ["المسارات", "99 إجمالاً، 73 في OpenAPI، و26 للواجهة"], ["فحص الترجمة", "ناجح"], ["فحص التغييرات", "ناجح"]], [3400, 5960], "الجدول 13: النتائج الحالية")
    add_heading(doc, "5.9 تفسير نتائج القبول المرحلي", 2)
    add_text(doc, "أغلقت Phase 5 بعد نجاح اختبار المختبر. أما Phase 6 فلها تنفيذ واختبارات، لكن ملف readiness يعلن النجاح، بينما يعلن تقرير القبول والاختبار الافتراضي أن البيئة محجوبة. ولا يوجد داخل المستودع سجل مستقل لقبول Phase 7 الحقيقي. لذلك نعرض هذه المعلومة كما هي ولا نعلن نجاحاً غير مثبت.")
    add_heading(doc, "5.10 قابلية الصيانة والتشغيل", 2)
    add_text(doc, "أصبح النظام أسهل للصيانة بسبب فصل العقود عن الخدمات والمستودعات والواجهات. ويمكن تغيير مزود النموذج أو واجهة العرض دون نقل صلاحية التنفيذ إليها. كما تساعد أدوات الجرد وملفات الرسوم ومصفوفة التتبع على فهم المشروع.")
    add_text(doc, "قبل التشغيل يفحص المشرف الإعدادات السرية واتصال قاعدة البيانات والمخطط وملف known_hosts. وبعد ذلك يفحص Ollama وMCP عند الحاجة. وتعرض لوحة الإدارة سجلات التنفيذ دون كلمات مرور أو رموز أو مفاتيح SSH.")
    add_heading(doc, "5.11 المشاكل والصعوبات والأعمال المؤجلة", 2)
    add_text(doc, "تشمل الأعمال المؤجلة الإشعارات الاجتماعية، والتنبؤ بالأعطال، والصيانة الاستباقية، ودراسة الاتجاهات الطويلة، والتعلم من قرارات المطور، والمقارنة مع OpenClaw، والاختبار الإنتاجي، وقبول clean-host، وتوحيد أدلة Phase 6 وPhase 7. وهي موثقة في docs/roadmap/deferred-requirements.md وfuture-work.md.")
    add_heading(doc, "5.12 تفاصيل تنفيذ المراقبة", 2)
    add_text(doc, "تستخدم خدمة المراقبة ملفاً معروفاً، ولا تقبل أمراً حراً من المستخدم أو النموذج. يحدد الملف الأوامر والمهل وحجم النتائج. وتحفظ كل نتيجة قبل تكوين التقرير، حتى نعرف سبب الفشل إذا توقف أمر أو انقطع الاتصال.")
    add_text(doc, "بعد ذلك ترتب الخدمة الحالات والأزمنة والمخرجات، وتحفظ التقرير في قاعدة البيانات. ويمكن للمشرف عرضه ومقارنة تقارير الخادم نفسه دون الاعتماد على ذاكرة جلسة Claude.")
    add_heading(doc, "5.13 تفاصيل تنفيذ التحليل وRAG", 2)
    add_text(doc, "يستخدم التحليل التقرير المحفوظ، ويضيف إليه الوثائق المناسبة عند الحاجة. ويميز النظام بين وثيقة عامة وتقرير حديث ونتيجة فحص متخصص. كما يحفظ مصدر كل معلومة حتى يستطيع المراجع فهم سبب ظهورها.")
    add_text(doc, "إذا تعطل Ollama أو لم يستطع النظام قراءة النتيجة، يحفظ سبب الفشل ولا ينفذ أي معالجة. ويبقى التقرير متاحاً للمراجعة أو لإعادة التحليل لاحقاً.")
    add_heading(doc, "5.14 تفاصيل تنفيذ التحقيق المتخصص", 2)
    add_text(doc, "تبدأ حلقة التحقيق بالمشكلة التي تحتاج إلى فحص. ثم تجلب المتخصصين الفعالين من قاعدة البيانات. لكل متخصّص هدف وأدوات وميزانية، وترفض الحلقة أي تعريف غير معروف أو متعارض.")
    add_text(doc, "بعد فحص السياسة تجمع الحلقة الأدلة بواسطة EvidenceCollectionService.collect. وما نحتاجه هو حفظ نتيجة الحلقة مع المصدر والمالك والوقت، ثم استخدام الدليل نفسه في التشخيص والعرض والتدقيق.")
    add_heading(doc, "5.15 تفاصيل تنفيذ المعالجة المشرفة", 2)
    add_text(doc, "تبدأ المعالجة المشرفة بخطة وإجراء واضحين، ثم تمر باختبار Sandbox أو بموافقة المستخدم حسب مستوى الخطر. وترتبط الموافقة بالخطة الحالية، لذلك لا تبقى صالحة إذا تغير الخادم أو الهدف أو الإجراء. وبعد التنفيذ تحفظ النتيجة والتحقق والتراجع.")
    add_text(doc, "لا يعني نجاح اتصال SSH أن المشكلة حُلّت. لذلك ينفذ النظام فحصاً مستقلاً بعد العملية. وإذا فشل الفحص، يسجل العملية كغير ناجحة ويبدأ الاستعادة عند الحاجة. وتظهر مراحل الموافقة والتنفيذ والتحقق في سجل الحادثة.")
    add_heading(doc, "5.16 تفاصيل تنفيذ المعالجة الذاتية والتعافي", 2)
    add_text(doc, "تحفظ قاعدة البيانات القرار والتصريح والحجز والتنفيذ والنتيجة في سجلات منفصلة. وبذلك نعرف هل سمحت السياسة، وهل استُهلك التصريح، وهل نفذ الإجراء، وهل نجح الفحص.")
    add_text(doc, "إذا انقطع العامل بعد الحجز، يفحص النظام حالة الحجز والتصريح وسجل التنفيذ قبل المتابعة. ويمكن لعامل جديد استرداد الحجز إذا انتهت ملكية العامل القديم. أما إذا كان التنفيذ غير واضح، فيتوقف النظام ويطلب مراجعة.")
    add_heading(doc, "5.17 دليل التشغيل والصيانة", 2)
    add_text(doc, "يحتاج تشغيل البيئة إلى إعداد قاعدة البيانات وملف known_hosts ومفاتيح SSH خارج المستودع. وبعد بدء الخدمة يفحص المشرف المخطط والأدوات وصحة الخدمات. وتبقى اختبارات القبول الحقيقي اختيارية حتى لا تتصل الاختبارات العادية بخادم حقيقي.")
    add_text(doc, "يراجع المشرف سجل التدقيق للبحث عن الفشل المتكرر أو انتهاء الحجوزات أو اختلاف الأدلة. وتساعد أدوات الجرد على تحديث عدد المسارات والأدوات والجداول. أما الوثائق القديمة فتظل مرجعاً تاريخياً حتى يتم تحديثها رسمياً.")


def chapter5_ui_and_testing(doc):
    add_heading(doc, "الفصل السادس: الواجهات والاختبار", 1)
    add_heading(doc, "6.1 مقدمة الواجهات", 2)
    add_text(doc, "تتكون واجهة المشروع من لوحة Admin مبنية على قوالب Jinja2 وملفات JavaScript وCSS. لا تعرض الواجهة كل تفاصيل قاعدة البيانات، بل تعرض المعلومات التي يحتاجها المشغل أو المدير لاتخاذ قرار ومتابعة حالته. وتستخدم الصفحات نفسها الخدمات التي تستخدمها API وأدوات MCP، لذلك لا يوجد تنفيذ مختلف مخفي داخل القالب.")
    add_heading(doc, "6.2 الواجهة الرئيسية ومتابعة الخوادم", 2)
    add_figure(doc, 25, "مخطط تخطيطي لواجهة لوحة الإدارة", "25-admin-dashboard.png")
    add_text(doc, "تعرض لوحة البداية حالة الخدمات والتقارير الحديثة والتحقيقات المفتوحة وآخر أحداث التدقيق. وتظهر الروابط الجانبية صفحات dashboard وservers وreports وinvestigations وaudit، بينما تبقى البيانات الحساسة خلف جلسة الإدارة والصلاحية المناسبة.")
    add_heading(doc, "6.3 واجهة التقارير والتحليل", 2)
    add_text(doc, "تتيح صفحة reports اختيار تقرير محفوظ وعرض تفاصيله ونتائج الأوامر ومصدر كل نتيجة. ومن صفحة التفاصيل يمكن طلب التحليل أو مراجعة تحليل سابق. عند نقص البيانات تظهر حالة غير مكتملة وسببها بدلاً من عرض نتيجة توحي بأن التقرير ناجح.")
    add_heading(doc, "6.4 واجهة التحقيق والأدلة", 2)
    add_figure(doc, 26, "مخطط تخطيطي لواجهة التحقيق والأدلة", "26-investigation-interface.png")
    add_text(doc, "تجمع صفحة investigations بين ملخص الحادثة ونتائج المتخصصين وEvidence ومصادرها. وتعرض الاختلافات بين المتخصصين صراحة، وتسمح للمشغل بالانتقال من التشخيص إلى الدليل الذي اعتمد عليه. أما صفحة specialists فتستخدم لإدارة التعريفات والأدوات والميزانيات، ولا تسمح بإضافة أداة غير مسجلة إلى مسار التحقيق.")
    add_heading(doc, "6.5 واجهة المعالجة والموافقة", 2)
    add_figure(doc, 27, "مخطط تخطيطي لواجهات المعالجة والسياسات", "27-remediation-interface.png")
    add_text(doc, "تعرض صفحة remediation الخطة وبصمتها والهدف ومستوى الخطورة ونتيجة الاختبار وحالة الموافقة. وتعرض صفحات policies وautonomous_authorizations وautonomous_reservations القرار والتصريح والحجز، بينما تعرض autonomous_history نتيجة التنفيذ والتحقق. لا تظهر owner token أو مفاتيح SSH أو الأسرار في أي صفحة.")
    add_heading(doc, "6.6 واجهة التدقيق وإدارة الجلسة", 2)
    add_text(doc, "تبدأ الإدارة من login، ثم تتحقق auth_routes من كلمة المرور والجلسة وتطبق routes الصلاحية على كل صفحة. تعرض audit هوية الفاعل والعملية والوقت والنتيجة وسبب الرفض عند توفره. ويعيد النظام المستخدم إلى صفحة الدخول عند انتهاء الجلسة، ويمنع الطلب المعدل عند غياب CSRF أو عند عدم كفاية الدور.")
    add_heading(doc, "6.7 ربط صفحات الواجهة بالوظائف", 2)
    add_table(doc, ["الصفحة أو القالب", "الوظيفة", "المعلومات المعروضة", "الصلاحية"], [
        ["dashboard.html", "ملخص حالة النظام", "الخدمات والتقارير والتحقيقات والأحداث الأخيرة", "viewer"],
        ["servers.html وmonitoring_profiles.html", "إدارة الخوادم والمراقبة", "الخادم والأوامر والمهلة والحالة", "operator"],
        ["reports.html وreport_details.html", "مراجعة التقرير والتحليل", "نتائج الأوامر والمصادر وحالة التحليل", "viewer/operator"],
        ["investigations.html وinvestigation_details.html", "التحقيق والأدلة", "المتخصصون والأدلة والتشخيص والاختلاف", "operator"],
        ["remediation.html", "مراجعة الخطة والتنفيذ", "البصمة والخطر والموافقة والتحقق", "operator/admin"],
        ["autonomous_policies.html وautonomous_reservations.html", "السياسة والحجز", "الحالة والمدة والمالك وسجل القرار", "admin"],
        ["audit.html", "المراجعة والتدقيق", "الفاعل والعملية والوقت والنتيجة", "viewer/admin"],
    ], [2600, 2500, 3300, 960], "الجدول 14: صفحات الواجهة ووظائفها")
    add_heading(doc, "6.8 خطة الاختبار", 2)
    add_text(doc, "اختبرت المنصة على طبقات تبدأ بالعقود والسياسات، ثم المستودعات، ثم خدمات المراقبة والتحقيق والمعالجة، ثم API وAdmin وMCP. وتستخدم الاختبارات البدائل المحلية للـSSH وOllama، بينما تبقى ملفات real_runtime اختيارية ولا تعمل ضمن regression العادي.")
    add_table(doc, ["نوع الاختبار", "ما الذي يفحصه", "النتيجة أو الحد"], [
        ["اختبار الوحدات", "قرارات السياسة وحدود الأوامر وتصنيف الحالات", "حتمي ولا يحتاج خدمات خارجية"],
        ["اختبار المستودعات", "الملكية والحجز والتصريح ومنع التكرار", "يفحص التزامن والدمج المشروط"],
        ["اختبار التكامل", "المراقبة والتحليل والتحقيق وربط Evidence", "يفحص انتقال الحالة بين الخدمات"],
        ["اختبار API وAdmin", "الجلسة وRBAC وCSRF والإسقاطات والقوالب", "لا يكشف الأسرار ولا يسمح بتجاوز backend"],
        ["اختبار MCP", "قائمة الأدوات والمدخلات والنتائج المنظمة", "25 أداة typed دون shell حر"],
        ["اختبار قاعدة البيانات", "الجداول والفهارس وpgvector والمهاجرات", "33 جدولاً و3 فهارس RAG"],
        ["القبول الحقيقي", "Ollama وPostgreSQL وSSH وSandbox وVPS", "opt-in ويتطلب بيئة وأسراراً وهدفاً آمناً"],
    ], [2200, 4700, 2460], "الجدول 15: مصفوفة أنواع الاختبار وحدودها")
    add_heading(doc, "6.9 نتائج الاختبار والقبول", 2)
    add_text(doc, "تسجل وثائق القبول النهائية 624 حالة في regression غير الحقيقي، نجحت منها 620 حالات، مع أربع حالات تخطٍّ متوقعة لبيئات real_runtime وعدم وجود فشل. وتبقى هذه النتيجة دليلاً على السلوك الحتمي المحلي، ولا تعني وحدها أن خدمات Ollama وSSH وSandbox متاحة في كل تشغيل.")
    add_text(doc, "أظهرت اختبارات الإدارة حماية الجلسات والصلاحيات وCSRF، وأظهرت اختبارات المعالجة منع إعادة استخدام التصريح ورفض الحفظ بعد فقدان الملكية، كما أثبتت اختبارات التحقيق ربط Evidence بالحادثة والخادم. وتعرض ملفات القبول الخارجي نتيجة كل مرحلة بصورة منفصلة عن هذه الأرقام.")


def appendices(doc):
    add_heading(doc, "الخاتمة والآفاق المستقبلية", 1)
    add_heading(doc, "1. مقدمة", 2)
    add_text(doc, "قدم هذا المشروع وكيلاً يساعد مشغل الخادم على جمع المعلومات وفهم الأعطال ومراجعة المعالجة قبل تنفيذها. بنيت مراحل العمل حول تقرير محفوظ، وتحليل مرتبط بمصدر، وتحقيق متخصص، ثم خطة وسياسة وتحقق وتدقيق. وبذلك أصبح النظام أقرب إلى أداة تشغيل يمكن مراجعتها، وليس مجرد محادثة تنتج اقتراحات.")
    add_heading(doc, "2. المشاكل والتحديات", 2)
    add_text(doc, "تمثلت الصعوبة الأولى في التعامل مع نتائج ناقصة أو قديمة، ولذلك ربط النظام كل نتيجة بمصدرها وزمنها وحادثتها. وتمثلت الصعوبة الثانية في التزامن؛ إذ يمكن لعاملين قراءة الحالة نفسها، لذلك يعتمد الحفظ النهائي على حجز قصير وowner token ودمج مشروط بدلاً من استبدال metadata كاملة من قراءة قديمة. كما احتاجت الإدارة إلى حماية مستقلة للجلسات وRBAC وCSRF، وإلى فصل واضح بين الاختبارات المحلية والقبول الحقيقي.")
    add_text(doc, "وتبقى بعض الحدود العملية قائمة: الإشعارات الاجتماعية وتحديد موضع خطأ الكود بصورة كاملة مؤجلان، وقبول البيئات الخارجية يعتمد على توفر Ollama وPostgreSQL وSSH وSandbox وهدف آمن. لا نعرض هذه الحدود على أنها وظائف منجزة، بل نثبتها كأعمال لاحقة أو شروط تشغيل.")
    add_heading(doc, "3. الآفاق المستقبلية", 2)
    for x in ["إضافة موصلات إشعار قابلة للتدقيق مع الحفاظ على سياسة الصلاحيات.", "توسيع تحليل الاتجاهات والتنبؤ بالأعطال بعد جمع تاريخ تشغيلي كافٍ.", "تحسين تحديد موضع خطأ الكود مع ربطه بملفات المصدر والاختبارات.", "إضافة إدارة متعددة المستأجرين وتقارير تشغيلية أطول مدى.", "توحيد قبول clean-host وPhase 6 وPhase 7 في سيناريوهات تشغيل قابلة لإعادة الإنتاج."]:
        add_bullet(doc, x)
    add_heading(doc, "4. الخاتمة", 2)
    add_text(doc, "يقدم المشروع بنية عملية لإدارة VPS بمساعدة الذكاء الصنعي مع إبقاء القرار والتنفيذ في خدمات Python وقاعدة PostgreSQL. تستخدم Claude للإشراف، وOllama للتحليل، وMCP لعرض القدرات المسموحة، بينما تجمع الحلقة المتخصصة Evidence فعلياً وتحفظه مخزن التحقيق المستمر. أثبتت الاختبارات المحلية حماية الواجهات والسياسات والتزامن ومنع التكرار، وتبقى الاختبارات الحقيقية مرتبطة ببيئة تشغيل محددة ومعلنة.")
    add_heading(doc, "المراجع والوثائق المستخدمة", 2)
    for x in ["docs/requirements/functional-requirements.md", "docs/requirements/non-functional-requirements.md", "docs/requirements/traceability-matrix.md", "docs/architecture/diagrams/ ومخططات النظام القابلة للتحرير", "docs/final-acceptance/ وتقارير القبول المرحلي", "docs/operations/ ودليل التشغيل والإعداد", "docs/testing/ ووثائق الاختبارات", "1-Report-Template.docx، واستخدمت التقارير السابقة لضبط أسلوب العرض فقط"]:
        add_bullet(doc, x)
    add_heading(doc, "الملحق أ: قائمة الجداول الكاملة", 1)
    add_table(doc, ["المجموعة", "الجداول"], [["المراقبة", "servers؛ monitor_commands؛ monitoring_profiles؛ monitoring_reports؛ command_executions"], ["التحليل والبحث", "report_analyses؛ knowledge_sources؛ knowledge_documents؛ knowledge_chunks"], ["التحقيق", "investigations؛ specialist_definitions؛ investigation_specialist_candidates؛ agent_jobs"], ["المعالجة المشرفة", "remediation_plans؛ approvals؛ executions؛ verifications؛ rollbacks؛ evidence؛ sandbox_validations"], ["التنفيذ الذاتي", "policies؛ decisions؛ authorizations؛ reservations؛ runtime_state؛ audit_events"], ["الإدارة", "admin_users؛ admin_sessions؛ admin_auth_audit_events"]], [2400, 6960], "الجدول 10: جرد الجداول وعددها 33")
    add_text(doc, "توضح القائمة وظيفة كل مجموعة من الجداول. تحفظ جداول المراقبة نتائج الخادم، وتحفظ جداول التحليل تفسير التقارير ومصادره. وتحفظ جداول التحقيق الأدلة والمتخصصين. أما جداول المعالجة فتحفظ الخطة والموافقة والتنفيذ والتحقق والتراجع، وتحتفظ جداول التنفيذ الذاتي بالقرار والتصريح والحجز وسجل التدقيق.")
    add_heading(doc, "الملحق أ-1: معنى مجموعات البيانات", 2)
    add_table(doc, ["المجموعة", "الغرض", "ضابط السلامة"], [["المراقبة", "حفظ تقارير الخادم ونتائج الأوامر", "الأوامر مسجلة والمخرجات محدودة"], ["التحليل والبحث", "تفسير التقارير وربط المصادر", "المصدر والهوية محفوظان"], ["التحقيق", "إدارة المتخصصين والأدلة", "الملكية وميزانية الأدوات"], ["المعالجة المشرفة", "المعالجة والموافقة والتحقق", "بصمة وخطر واستعادة"], ["التنفيذ الذاتي", "التنفيذ المقيد والتعافي", "تصريح واحد وحجز ومالك"], ["الإدارة", "الهوية والجلسات والتدقيق", "scrypt وCSRF وRBAC"]], [2300, 4300, 2760], "الجدول 11: دلالة مجموعات قاعدة البيانات")
    add_heading(doc, "الملحق ب: مسارات الواجهات", 1)
    add_text(doc, "يحتوي المشروع على 99 مساراً، منها 73 مساراً ظاهراً في OpenAPI و26 مساراً خاصاً بالواجهة. وتشمل المسارات الجديدة صفحات السياسات والحالات والحجوزات والتفويضات والتدقيق، إضافة إلى مسار سجل مصادقة الإدارة.")
    add_text(doc, "تختبر مسارات API العقود والاستجابات والصلاحيات، بينما تختبر مسارات الواجهة الجلسة وCSRF والقوالب. ولا يعني وجود مسار جديد أنه يملك صلاحية جديدة، لأن جميع المسارات تمر بخدمات الخادم التي تطبق القيود.")
    add_heading(doc, "الملحق ب-1: عينة حالات الاستخدام", 2)
    add_table(doc, ["المعرف", "الحالة", "الممثل", "النتيجة"], [["UC-001", "تشغيل مراقبة خادم", "المشرف وClaude", "تقرير محفوظ مع نتائج الأوامر"], ["UC-005", "تحليل تقرير سابق", "Python وOllama", "تحليل مربوط بمصادره"], ["UC-009", "طلب تحقيق متخصص", "موجّه التحقيق", "أدلة وتشخيص مستمران"], ["UC-013", "إنشاء خطة معالجة", "خدمة المعالجة", "خطة ذات بصمة ومخاطر"], ["UC-017", "موافقة بشرية", "المشغل والمدير", "موافقة مرتبطة بالخطة"], ["UC-019", "تنفيذ ذاتي آمن", "عامل التنفيذ", "تنفيذ واحد أو رفض"], ["UC-022", "استعادة حجز", "عامل التعافي", "استئناف مشروط أو توقف آمن"], ["UC-025", "إدارة لوحة النظام", "المدير", "إدارة دون كشف أسرار"]], [1400, 3100, 2500, 2360], "الجدول 12: عينة حالات الاستخدام الأساسية")
    add_heading(doc, "الملحق ج: مصادر التحقق", 1)
    for x in ["المصدر البرمجي: app/core، app/capabilities، app/runtime/claude، app/interfaces، app/infrastructure، app/composition.", "الاختبارات: tests/، مع استبعاد tests/real_runtime من regression العادي.", "التحقق: tools/bootstrap_database.py، tools/dev/list_routes.py، ProjectMcpToolBoundary.list_tools().", "الوثائق المرجعية: دفتر الشروط PDF، القالب 1-Report-Template.docx، والتقرير العربي القديم كمصدر مصطلحات فقط.", "الهوية الطلابية الموثقة في هذه النسخة: سعيد بقدونس."]:
        add_bullet(doc, x)
    add_heading(doc, "الملحق د: مصفوفة الاختبارات", 1)
    add_table(doc, ["المجال", "ما الذي يثبت", "النتيجة الحالية"], [["العقود والسياسات", "حالات القرار والحدود والربط", "مغطى في الاختبارات"], ["المستودعات", "الملكية وعدم التكرار والتزامن", "مغطى في اختبارات التزامن"], ["الأمن", "الصلاحيات وCSRF وعدم كشف الأسرار", "ناجح"], ["MCP", "قائمة الأدوات والمدخلات المنظمة", "25 أداة"], ["قاعدة البيانات", "الجداول والفهارس وpgvector", "33 من 33 و3 فهارس"], ["لوحة الإدارة", "الجلسات والصفحات والبيانات المعروضة", "مغطى في TestClient"], ["القبول الحقيقي", "Ollama وSSH وSandbox الحقيقية", "اختياري وغير مشغل"]], [2300, 4700, 2360], "الجدول 13: مصفوفة الاختبارات ونتائجها")
    add_text(doc, "تُقرأ نتيجة 586 اختباراً مع هذا التفصيل: هي دليل على سلامة السلوك الحتمي في البيئة الآمنة، وليست بديلاً عن قبول خارجي يتطلب خدمات حقيقية. ولذلك يحافظ التقرير على فصل واضح بين ما يمكن إعادته محلياً وما يحتاج بنية تشغيلية وأسراراً وهدفاً آمناً.")
    add_heading(doc, "الملحق هـ: قاموس المكونات", 1)
    for x in ["app/core: العقود والإعدادات والسياسات والحالات الأساسية.", "app/capabilities: قدرات المراقبة والتحليل والتحقيق والمعالجة، وهي موضع السلطة التنفيذية.", "app/infrastructure: PostgreSQL والنماذج والمستودعات وSSH وOllama.", "app/interfaces: واجهات MCP وAdmin وAPI، مع بقاء backend authority.", "app/runtime/claude: عقد التشغيل الإشرافي وتدفق الحوار المنظم.", "app/composition: تركيب المستودعات والخدمات والruntime في حاوية واحدة.", "tests: اختبارات حتمية واختبارات opt-in للبيئات الخارجية.", "tools/dev: أدوات جرد المسارات وتزامن الوثائق والتحقق وبناء التقرير."]:
        add_bullet(doc, x)


def _replace_generated_paragraph(paragraph, text):
    """Replace one generated paragraph while retaining the report's RTL layout."""
    paragraph.text = text
    set_rtl(paragraph)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.15
    for run in paragraph.runs:
        set_font(run, size=12)


def synchronize_final_facts(doc):
    """Synchronize late acceptance facts without changing production code."""
    replacements = {
        "تم تنفيذ الوظائف": "ينفذ المشروع الوظائف الأساسية، بما فيها تصنيف SPEC-03 بثلاث حالات مستقلة: normal وdangerous وsensitive. ولا يجوز خلط هذا التصنيف مع شدة التحليل info وwarning وcritical أو مع خطر المعالجة low وmedium وhigh وcritical؛ إذ تمنع السياسة التنفيذ الذاتي للحالتين dangerous وsensitive. كما ينفذ SPEC-07 تحديداً منظماً ومحدوداً لموضع خطأ التطبيق من traceback أو السجلات، مع المسار ورقم السطر عند توفره واسم الدالة أو الوحدة عند توفره، وربط السبب والاستثناء بـ Evidence وحفظه وإظهاره في التحقيق والتشخيص النهائي وواجهات API وAdmin. أما SPEC-08 فتبقى فاشلة حرفياً ومقبولة كاستثناء مشروع: لا توجد قناة Telegram أو قناة اجتماعية، بينما تمثل موافقة Admin سير عمل الموافقة البشرية المنفذ. لا يشكل هذا الاستثناء عائقاً لإغلاق المشروع.",
        "توجد المتطلبات غير الوظيفية": "توجد المتطلبات غير الوظيفية في docs/requirements/non-functional-requirements.md. ومن أمثلتها حماية الجلسات وCSRF، منع الأدوات الحرة، اكتمال 33 من 33 جدولاً، وجود pgvector و3 من 3 فهارس RAG، توفير 25 أداة MCP، ومنع تنفيذ العملية نفسها أكثر من مرة. وفي الانحدار الحتمي النهائي جُمعت 624 حالة، نجحت 620، وتخطت 4 حالات اختيارية، ولم تفشل أي حالة، مع تحذير Starlette/httpx موجود مسبقاً وزمن 30.00 ثانية.",
        "تبدأ المعالجة المشرفة": "تبدأ المعالجة المشرفة بخطة وإجراء مسجلين، ثم تُمارس العملية المسجلة في مسار التحقق داخل Sandbox الأصلي المعزول، لا كفحص metadata أو dry-run فقط. يحفظ المسار Evidence قبل التنفيذ، وينفذ الإجراء، ويحفظ Evidence بعده، ويتحقق من الحالة المتوقعة، ثم ينفذ الإجراء العكسي ويحفظ Evidence الاستعادة ويتأكد من رجوع الحالة الأصلية. وترتبط كل نتيجة ببصمة الخطة، وتحمي من الخطة القديمة، وتفشل مغلقاً عند نقص البيانات أو اختلافها. كما ترتبط الموافقة بالخطة الحالية، لذلك لا تبقى صالحة إذا تغير الخادم أو الهدف أو الإجراء.",
        "يعتمد Admin": "يعتمد Admin على جلسات server-side، وتخزين كلمات المرور باستخدام scrypt، وحماية CSRF، وأدوار viewer وoperator وadmin ضمن نموذج صلاحيات مركزي. تدعم الجلسات الإبطال، ويسجل التدقيق هوية الفاعل، وتخضع نقاط API وWeb إلى middleware مشترك حتى لا تصبح إخفاءات الواجهة بديلاً عن الحماية. وتستخدم إسقاطات العرض قائمة حقول صريحة، ولذلك لا تصل رموز التفويض أو owner token أو المفاتيح الخاصة إلى المتصفح.",
        "تغطي الاختبارات": "تغطي الاختبارات العقود والسياسات والمستودعات والواجهات والأمان. وتفحص اختبارات الوحدات قرارات السياسة وحدود الأوامر وتجميع الأدلة. كما تفحص اختبارات المستودعات التزامن والحجز والتصريح ومنع التكرار، وتثبت اختبارات التحقيق استخراج موضع خطأ التطبيق من traceback أو log وربطه وحفظه وإظهاره عبر API وAdmin.",
        "أغلقت Phase 5": "أغلقت Phase 5 بعد نجاح اختبار المختبر. أما القبول الحقيقي للتشغيل الإشرافي، فقد نجح مسار Claude Specialist ووصل إلى التنفيذ القانوني وحفظ Evidence في قاعدة المشروع. ولم يثبت إتمام جميع المتخصصين وإنهاء التحقيق داخل مهلة القبول الثابتة البالغة 300 ثانية؛ لذلك يصنف هذا الحد NONDETERMINISTIC_SUPERVISORY_ACCEPTANCE، مع ACCEPTED_LIMITATION = YES وPROJECT_CLOSURE_BLOCKING = NO. المسارات الأساسية للتنفيذ والحفظ وEvidence مثبتة بصورة مستقلة، ولم يحدد القبول عيباً في orchestration الإنتاجي أو يبرر تغييراً معمارياً.",
        "يقدم المشروع طريقة آمنة": "يقدم المشروع طريقة آمنة لمراقبة الخوادم وتحليل أعطالها. يستخدم Claude للإشراف، وOllama للمساعدة في التحليل، وMCP لعرض 25 أداة مسموحة، بينما تنفذ Python عمليات الفحص والحفظ والتدقيق. أثبتت الاختبارات حماية الصلاحيات وCSRF، وحدود الأدوات، وحفظ الأدلة، ومنع التكرار، والتعامل مع الاستعادة. كما تحققنا من 33 جدولاً وpgvector و3 فهارس RAG و9 تعريفات Specialist. وتبقى SPEC-08 فاشلة حرفياً مع استثناء مقبول وغير حاجب للإغلاق، بينما تبقى حدود القبول الإشرافي الحقيقي موثقة كقيد غير حتمي لا كعيب منتج.",
        "تُقرأ نتيجة 586": "تُقرأ نتيجة الانحدار النهائي 624 حالة: 620 ناجحة و4 متخطاة اختيارياً و0 فشل، بزمن 30.00 ثانية وتحذير Starlette/httpx موجود مسبقاً. وهي دليل على سلامة السلوك الحتمي في البيئة الآمنة، وليست بديلاً عن القبول الخارجي الذي يتطلب خدمات حقيقية. ولذلك يحافظ التقرير على فصل واضح بين ما يمكن إعادته محلياً وما يحتاج بنية تشغيلية وأسراراً وهدفاً آمناً.",
    }
    for paragraph in doc.paragraphs:
        for prefix, replacement in replacements.items():
            if paragraph.text.startswith(prefix):
                _replace_generated_paragraph(paragraph, replacement)
                break
    caption_numbers = {
        "الجدول 10: جرد الجداول": "الجدول 16: جرد الجداول",
        "الجدول 11: دلالة مجموعات قاعدة البيانات": "الجدول 17: دلالة مجموعات قاعدة البيانات",
        "الجدول 12: عينة حالات الاستخدام الأساسية": "الجدول 18: عينة حالات الاستخدام الأساسية",
        "الجدول 13: مصفوفة الاختبارات ونتائجها": "الجدول 19: مصفوفة الاختبارات ونتائجها",
    }
    for paragraph in doc.paragraphs:
        if paragraph.text in caption_numbers:
            _replace_generated_paragraph(paragraph, caption_numbers[paragraph.text])
    for table in doc.tables:
        for row in table.rows:
            labels = [cell.text for cell in row.cells]
            if labels and "تحديد موضع خطأ الكود" in labels[0]:
                row.cells[1].text = "منفذة ضمن حدود منظمة"
                row.cells[2].text = "traceback/log؛ المسار والسطر والدالة؛ Evidence والتشخيص وAPI/Admin"
            elif labels and "بيئة الاختبار المعزولة" in labels[0]:
                row.cells[1].text = "منفذة في المسار المخصص"
                row.cells[2].text = "قبل/بعد Evidence؛ تحقق؛ عكس؛ استعادة؛ بصمة وفشل مغلق"
            elif labels and "الإشعار الاجتماعي" in labels[0]:
                row.cells[1].text = "فشل مقبول كاستثناء مشروع"
                row.cells[2].text = "لا Telegram أو قناة اجتماعية؛ الموافقة الإدارية هي البديل البشري"
            for cell in row.cells:
                if "586 ناجحاً" in cell.text:
                    cell.text = "624 مجموعة؛ 620 ناجحة؛ 4 متخطاة؛ 0 فشل؛ تحذير واحد؛ 30.00 ثانية"
                    for paragraph in cell.paragraphs:
                        set_rtl(paragraph)
                        for run in paragraph.runs:
                            set_font(run, size=8.5)


def main() -> int:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document(str(TEMPLATE))
    clear_body(doc)
    configure_document(doc)
    cover(doc)
    abstract_and_front(doc)
    add_text(doc, "ملاحظة: يعتمد هذا التقرير على حالة المشروع الحالية. استخدمنا التقارير السابقة لتنسيق الأسلوب فقط، ولم ننقل منها معلومات عن المشروع.", bold=True, size=11)
    add_heading(doc, "مقدمة عامة", 1)
    add_text(doc, "تحتاج إدارة الخوادم إلى متابعة مستمرة للموارد والخدمات والسجلات. وقد يؤدي تأخر اكتشاف المشكلة إلى توقف الخدمة أو فقدان البيانات. لذلك أصبح من المفيد استخدام أدوات تساعد المشرف في جمع المعلومات وتحليلها، مع إبقاء القرار النهائي تحت السيطرة.")
    add_text(doc, "يعرض هذا المشروع وكيلاً ذكياً يساعد في مراقبة خوادم Linux وفهم الأعطال التي تظهر فيها. يجمع الوكيل المعلومات من أوامر محددة، ويحفظها في قاعدة البيانات، ثم يستخدم التحليل والأدلة للوصول إلى نتيجة مفهومة. وإذا احتاجت المشكلة إلى معالجة، يمررها النظام عبر خطوات تحقق وموافقة قبل التنفيذ.")
    chapter1(doc)
    chapter2(doc)
    chapter3(doc)
    chapter4(doc)
    chapter5_ui_and_testing(doc)
    appendices(doc)
    synchronize_final_facts(doc)
    doc.core_properties.title = "التقرير التقني لوكيل ذكاء صنعي آمن ومستقل لإدارة الخوادم الافتراضية الخاصة"
    doc.core_properties.subject = "Implementation and testing report"
    doc.core_properties.author = "سعيد بقدونس"
    doc.core_properties.comments = "Generated from 1-Report-Template.docx; Arabic implementation-focused technical report."
    doc.save(str(OUT))
    print(str(OUT).encode("ascii", "backslashreplace").decode())
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
